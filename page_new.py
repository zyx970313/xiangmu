from flask import Flask, render_template, request, flash, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
import requests
from bs4 import BeautifulSoup
import jieba
from KB_query.word_tagging import Tagger
import KB_query.jena_sparql_endpoint
import KB_query.question2sparql
import xlrd
import time
import pymysql
import pandas as pd
from KB_query.jena_sparql_endpoint import JenaFuseki
from KB_query.question2sparql import Question2Sparql
from sklearn.feature_extraction.text import CountVectorizer
from scipy.linalg import norm,pinv
import math
import re
from sklearn.cluster import AffinityPropagation
import nltk
from sklearn.decomposition import PCA
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from selenium import webdriver
import pywinauto
from pywinauto.keyboard import send_keys
from selenium.webdriver import Chrome, ChromeOptions
from docx import Document
import os
from urllib.request import  urlopen
from sklearn import feature_extraction
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.feature_extraction.text import CountVectorizer
from tf import tfi,tfi1
from sklearn.cluster import KMeans
import cv2
import json
import argparse
import imutils
import torch
from io import BytesIO
from crnn import LiteCrnn, CRNNHandle
from psenet import PSENet, PSENetHandel
from application import idcard, trainTicket
from crnn.keys import alphabetChinese as alphabet
from angle_class import AangleClassHandle, shufflenet_v2_x0_5
from scipy.spatial import distance as dist
from collections import OrderedDict
from PIL import Image
import sys
import xlsxwriter  # 导入模板
#切图
def cut_image(image):
    width, height = image.size
    item_height = int(height / 10)
    box_list = []
    box=(0,0,width,item_height)
    box_list.append(box)
    box=(0,item_height,width,height)
    box_list.append(box)
    image_list = [image.crop(box) for box in box_list]
    return image_list

#保存
def save_images(image_list):
    index = 1
    for image in image_list:
        image.save('./output'+str(index) + '.png', 'PNG')
        index += 1

def upload_files(a):
    url = "http://localhost:3030/dataset.html?tab=upload&ds=/faults_last.html"
    opt = ChromeOptions()  # 创建Chrome参数对象
    opt.headless = True  # 把Chrome设置成可视化无界面模式
    browser = Chrome(options=opt)  # 创建Chrome无界面对象
    # 访问图片上传的网页地址
    browser.get(url=url)
    time.sleep(3)
    # 点击图片上传按钮，打开文件选择窗口
    browser.find_element_by_name("files[]").send_keys(r"D:\pycode\library\d2rq\d2rq-0.8.1\faults_last.nt")
    button = browser.find_element_by_css_selector("[class='btn btn-primary start action-upload-all']")
    button.click()
    browser.close()
    print(a)

    return url

def create_nt(a):
    command1 = "cd /d D:\pycode\library\d2rq\d2rq-0.8.1"
#    command2 = "generate-mapping -u root -p zhangyixin -o faults3.ttl jdbc:mysql:///faults2"
    command3 = r".\dump-rdf.bat -o faults_last.nt .\faults_last.ttl"
#    cmd = "{0} && {1} && {2}".format(command1, command2, command3)
    cmd = "{0} && {1}".format(command1, command3)
    os.system(cmd)
    print(a)
    return command1


app = Flask(__name__)
app.secret_key = 'test'
# 设置数据库的连接地址
app.config["SQLALCHEMY_DATABASE_URI"] = "mysql://root:zhangyixin@127.0.0.1:3306/faults_last"
# 是否监听数据库变化  一般不打开, 比较消耗性能
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
# 创建数据库操作对象(建立了数据库连接)
db = SQLAlchemy(app)

app.config["SECRET_KEY"] = "zzzyyyxxx"

# 故障特征表
class Feature(db.Model):
    __tablename__ = "features"
    features_id = db.Column(db.Integer, primary_key=True)
    features = db.Column(db.String(64))
    group_id = db.Column(db.Integer)
# 关系表
class Fea2Sol(db.Model):
    __tablename__ = "fea2sol"
    features_id = db.Column(db.Integer, db.ForeignKey("features.features_id"), primary_key=True)
    sol_features_id = db.Column(db.Integer, db.ForeignKey("sol_features.sol_features_id"),primary_key=True)

# 解决方案特征表
class Sol_feature(db.Model):
    __tablename__ = "sol_features"
    sol_features_id = db.Column(db.Integer, primary_key=True)
    sol_features = db.Column(db.String(64))
    group_id = db.Column(db.Integer)


class RBF:
    def __init__(self,input_dim,num_centers,out_dim,centers):
        self.input_dim=input_dim
        self.num_centers=num_centers
        self.out_dim=out_dim
        self.beta=2
        self.centers=centers
        self.W=[1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1]


    def _basisfunc(self,c,d):
        return np.exp(-self.beta* norm(c-d)**2)

    def _cacAct(self,X):
        G=np.zeros((X.shape[0],self.num_centers),dtype=np.float)
        for ci,c in enumerate(self.centers):
            for xi,x in enumerate(X):
                G[xi,ci]=self._basisfunc(c,x)
        return G


    def train(self,X):
     #   rnd_idx=np.random.permutation(X.shape[0])[:self.num_centers]
        self.centers=self.centers
        #RBF的激活函数的值
        G=self._cacAct(X)
        print(G)
        print('G[0]',G[0])
        print('G[0][0]',G[0][0])
        global after_sort, after_sort_index, similar_list
        similar_list = G[0]

        global max_G,max_G_id
        max_G = 0
        max_G_id = 0
        for i in range(0,len(G[0])-1):
            if G[0][i]>max_G:
                max_G_id=i
                max_G=G[0][i]

        print('max_G_id',max_G_id)
        print('max_G',max_G)

        #self.W=np.dot(pinv(G),Y)
        #print(self.W)


    # def predict(self,X):
    #     G=self._cacAct(X)
    #     Y=np.dot(G,self.W)
    #     return Y


class Vertex(object):
	# 初始化顶点
    def __init__(self, key, text):
        self.text = text
        self.id = key  # 初始化顶点的键
        self.connectedTo = {}


# 添加邻居顶点，参数nbr是邻居顶点的键，默认权重为0
    def addNeighbor(self, nbr, weight=0):
         self.connectedTo[nbr] = weight

    def __str__(self):
        return str(self.id) + ' connectedTo: ' + str([x.id for x in self.connectedTo])

    # 获取该顶点所有邻居顶点的键
    def getConnections(self):
        return self.connectedTo.keys()

    # 获取顶点的键
    def getId(self):
        return self.id

    # 获取到某邻居顶点的权重
    def getWeight(self, nbr):
        return self.connectedTo[nbr]

    #获取节点文字
    def getText(self):
        return self.text


# 自定义图类
class Graph(object):
    # 初始化图
    def __init__(self):
        self.vertList = {}
        self.numVertices = 0


    # 添加顶点
    def addVertex(self, key,text):
        newVertex = Vertex(key,text)					#创建顶点
        self.vertList[key] = newVertex 			#将新顶点添加到邻接表中
        self.numVertices = self.numVertices + 1 #邻接表中顶点数+1
        return newVertex

    # 获取顶点
    def getVertex(self, n):
        if n in self.vertList:					#若待查询顶点在邻接表中，则
            return self.vertList[n] 			#返回该顶点
        else:
            return None

    # 使之可用in方法
    def __contains__(self, n):
        return n in self.vertList

    # 添加边，参数f为起始顶点的键，t为目标顶点的键，cost为权重
    def addEdge(self, f, t, cost=0):
        if f not in self.vertList:				#起始顶点不在邻接表中，则
            self.addVertex(f) 					#添加起始顶点
        if t not in self.vertList:				#目标顶点不在邻接表中，则
            self.addVertex(t)					#添加目标顶点
        self.vertList[f].addNeighbor(self.vertList[t], cost)#在邻接表中添加起始点的目标点及权重

    # 获取邻接表中所有顶点的键
    def getVertices(self):
        return self.vertList.keys()

    # 迭代显示邻接表的每个顶点的邻居节点
    def __iter__(self):
        return iter(self.vertList.values())


def findPath(graph,start,end,path=[]):
    path = path + [start]
    if start == end:
        return path

    paths = []  # 存储所有路径
    for node in graph[start]:
        if node not in path:
            newpaths = findPath(graph, node, end, path)
            for newpath in newpaths:
                paths.append(newpath)
    return paths


#精确查询中的特征及解决办法
y1,y2,y3,y4,y5,y6,y7,y8,y9,y10=None,None,None,None,None,None,None,None,None,None
w1,w2,w3,w4,w5,w6,w7,w8=None,None,None,None,None,None,None,None
s1,s2,s3,s4,s5,s6,s7,s8,s9,s10="","","","","","","","","",""
#精确查询中每个特征在数据库中的数量
num1,num2,num3,num4,num5,num6,num7,num8,num9,num10,num11,num12,num13,num14,num15,num16,num17,num18="","","","","","","","","","","","","","","","","",""
#模糊查询中的三个案例的id及特征

f1,f2,f3,f4,f5,f6,f7,f8,f9,f10=None,None,None,None,None,None,None,None,None,None
#批量上传数据
list_fs=[]
#模糊查询
fuzz_q1=[]
fuzz1=[]

last_fea_count=0
last_sol_count=0
next_id=0

max_G=0
max_G_id=0
after_sort=[]
after_sort_dict_index=[]
similar_list=[]

phenomenon_six = ["", "", "", "", "", ""]
method_six = ["", "", "", "", "", ""]
case_num=0

list_chart,list_chart1 = [],[]
links,links1 = [],[]
case_num=0
allpath=[]
text_concat=[]
case_text=[]
v1=0
desc_text=[]
prec_text=[]

list_map=[]
links_map=[]

def open_demo(image):
    kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (17, 17))
    binary = cv2.morphologyEx(image, cv2.MORPH_OPEN, kernel)  # 开操作
    cv2.imshow("open-inv", binary)
    return binary

def close_demo(image):
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (28, 28))
    binary = cv2.morphologyEx(image, cv2.MORPH_CLOSE, kernel)  # 闭操作
    cv2.imshow("close_demo", binary)
    return binary

def crop_rect(img, rect, alph=0.15):
    img = np.asarray(img)
    center, size, angle = rect[0], rect[1], rect[2]
    min_size = min(size)
    if (angle > -45):
        center, size = tuple(map(int, center)), tuple(map(int, size))
        size = (int(size[0] + min_size * alph), int(size[1] + min_size * alph))
        height, width = img.shape[0], img.shape[1]
        M = cv2.getRotationMatrix2D(center, angle, 1)
        img_rot = cv2.warpAffine(img, M, (width, height))
        img_crop = cv2.getRectSubPix(img_rot, size, center)
    else:
        center = tuple(map(int, center))
        size = tuple([int(rect[1][1]), int(rect[1][0])])
        size = (int(size[0] + min_size * alph), int(size[1] + min_size * alph))
        angle -= 270
        height, width = img.shape[0], img.shape[1]
        M = cv2.getRotationMatrix2D(center, angle, 1)
        img_rot = cv2.warpAffine(img, M, (width, height))
        img_crop = cv2.getRectSubPix(img_rot, size, center)
    img_crop = Image.fromarray(img_crop)
    return img_crop

def crnnRec(im, rects_re, f=1.0):
    results = []
    im = Image.fromarray(im)
    for index, rect in enumerate(rects_re):
        degree, w, h, cx, cy = rect
        partImg = crop_rect(im, ((cx, cy), (h, w), degree))
        newW, newH = partImg.size
        partImg_array = np.uint8(partImg)
        if newH > 1.5 * newW:
            partImg_array = np.rot90(partImg_array, 1)
        angel_index = angle_handle.predict(partImg_array)
        angel_class = lable_map_dict[angel_index]
        rotate_angle = rotae_map_dict[angel_class]
        if rotate_angle != 0:
            partImg_array = np.rot90(partImg_array, rotate_angle // 90)
        partImg = Image.fromarray(partImg_array).convert("RGB")
        partImg_ = partImg.convert('L')
        try:
            if crnn_vertical_handle is not None and angel_class in ["shudao", "shuzhen"]:
                simPred = crnn_vertical_handle.predict(partImg_)
            else:
                simPred = crnn_handle.predict(partImg_)  # 识别的文本
        except:
            continue
        if simPred.strip() != u'':
            results.append({'cx': cx * f, 'cy': cy * f, 'text': simPred, 'w': newW * f, 'h': newH * f,
                            'degree': degree})
    return results


def text_predict(img):
    '''文本预测'''
    preds, boxes_list, rects_re, t = text_handle.predict(img, long_size=pse_long_size)
    result = crnnRec(np.array(img), rects_re)
    return result

def tiqu(s):
    two = []
    pattern = re.compile('发射车|装填车|装运车|运输车')
    car = pattern.search(s)
    if car is not None:
        print('car', car.group())
        s = s.replace(car.group(), '')
        two.append(car.group())
    else:
        two.append(' ')
    pattern1 = re.compile('1号车|2号车|3号车|4号车|5号车|6号车|7号车|8号车|9号车|10号车')
    type = pattern1.search(s)
    if type is not None:
        print('type', type.group())
        s = s.replace(type.group(), '')
        two.append(type.group())
    else:
        two.append(' ')
    print('two', two)
    weather = []
    temperature = re.findall(r'-?\d+\.?\d*e?-?\d*?℃', s)
    if len(temperature) == 0:
        two.append(' ')
        two.append(' ')
    elif len(temperature) == 1:
        two.append(' ')
    else:
        two.append(temperature[0])
        two.append(temperature[1])
    humidity = re.findall(r'-?\d+\.?\d*e?-?\d*%', s)
    if len(humidity) == 0:
        two.append(' ')
    else:
        two.append(humidity[0])
    pressure = re.findall(r'-?\d+\.?\d*e?-?\d*hpa', s)
    if len(pressure) == 0:
        two.append(' ')
    else:
        two.append(pressure[0])
    height = re.findall(r'-?\d+\.?\d*e?-?\d*m', s)
    if len(height) == 0:
        two.append(' ')
    else:
        two.append(height[0])
    salinity = re.findall(r'-?\d+\.?\d*e?-?\d*ppm', s)
    if len(salinity) == 0:
        two.append(' ')
    else:
        two.append(salinity[0])
    # salinity=salinity.replace('ppm',' ')
    print('temperature', temperature)
    print('humidity', humidity)
    print('pressure', pressure)
    print('height', height)
    print('salinity', salinity)
    print('two', two)

    return two

# 调用CPU或GPU
gpu_id = 0
if gpu_id and isinstance(gpu_id, int) and torch.cuda.is_available():
    device = torch.device("cuda:{}".format(gpu_id))
else:
    device = torch.device("cpu")
print('device:', device)

# psenet相关
pse_scale = 1
pse_long_size = 960  # 图片长边
pse_model_type = "mobilenetv2"
pse_model_path = "models/psenet_lite_mbv2.pth"
text_detect_net = PSENet(backbone=pse_model_type, pretrained=False, result_num=6, scale=pse_scale)
text_handle = PSENetHandel(pse_model_path, text_detect_net, pse_scale, gpu_id=gpu_id)

# crnn相关
nh = 256
crnn_model_path = "models/crnn_lite_lstm_dw_v2.pth"
crnn_net = LiteCrnn(32, 1, len(alphabet) + 1, nh, n_rnn=2, leakyRelu=False, lstmFlag=True)
crnn_handle = CRNNHandle(crnn_model_path, crnn_net, gpu_id=gpu_id)
crnn_vertical_model_path = "models/crnn_dw_lstm_vertical.pth"
crnn_vertical_net = LiteCrnn(32, 1, len(alphabet) + 1, nh, n_rnn=2, leakyRelu=False, lstmFlag=True)
crnn_vertical_handle = CRNNHandle(crnn_vertical_model_path, crnn_vertical_net, gpu_id=gpu_id)

# angle_class相关
lable_map_dict = {0: "hengdao", 1: "hengzhen", 2: "shudao", 3: "shuzhen"}  # hengdao: 文本行横向倒立 其他类似
rotae_map_dict = {"hengdao": 180, "hengzhen": 0, "shudao": 180, "shuzhen": 0}  # 文本行需要旋转的角度
angle_model_path = "models/shufflenetv2_05.pth"
angle_net = shufflenet_v2_x0_5(num_classes=len(lable_map_dict), pretrained=False)
angle_handle = AangleClassHandle(angle_model_path, angle_net, gpu_id=gpu_id)

def result(img):
    back = {}
    img = Image.open(BytesIO(img)).convert('RGB')
    img = np.array(img)
    result = text_predict(img)
    back['文本'] = list(map(lambda x: x['text'], result))
    res = trainTicket.trainTicket(result)
    back['火车票'] = str(res)
    res = idcard.idcard(result)
    back['身份证'] = str(res)
    return back


def temp_init():
    global y1, y2, y3, y4, y5, y6,y7,y8,y9,y10,s1, s2, s3, s4, s5, s6,s7,s8,s9,s10,num1, num2, num3, num4, num5, num6,num7,num8,num9,num10
    global f1, f2, f3, f4, f5, f6,f7,f8,f9,f10,list_fs,fuzz_q1,fuzz_q2,fuzz_q3
    global fuzz1,fuzz2,fuzz3
    # 精确查询中的特征及解决办法
    y1, y2, y3, y4, y5, y6,y7,y8,y9,y10 = None, None, None, None, None, None, None, None, None, None
    s1, s2, s3, s4, s5, s6,s7,s8,s9,s10 = "", "", "", "", "", "","","","",""
    # 精确查询中每个特征在数据库中的数量
    num1, num2, num3, num4, num5, num6,num7,num8,num9,num10 = "", "", "", "", "", "","","","",""
    # 模糊查询中的三个案例的id及特征

    f1, f2, f3, f4, f5, f6,f7,f8,f9,f10 = None, None, None, None, None, None, None, None, None, None
    # 批量上传数据
    list_fs = []
    # 模糊查询
    fuzz_q1 = ['', '', '', '', '', '']
    fuzz_q2 = ['', '', '', '', '', '']
    fuzz_q3 = ['', '', '', '', '', '']
    fuzz1 = ['', '', '', '', '', '']
    fuzz2 = ['', '', '', '', '', '']
    fuzz3 = ['', '', '', '', '', '']
    return 0


#主页
@app.route("/", methods=['GET', 'POST'])
def index():
    temp_init()
    if request.method=='POST':
        uname=request.form.get("uname")
        upwd=request.form.get("upwd")
        print("uname",uname)
        print("upwd",upwd)
        connection = pymysql.connect(host='localhost',
                                     user='root',
                                     passwd='zhangyixin',
                                     db='faults_last',
                                     port=3306,
                                     charset='utf8'
                                     )
        cur = connection.cursor()  # 游标（指针）cursor的方式操作数据
        sql1 = 'SELECT * FROM users where username="'+uname+'" and password="'+upwd+'"'
        try:
            cur.execute(sql1)  # execute(query, args):执行单条sql语句。
            see1 = cur.fetchone()  # 使结果全部可看
            print(see1[0])
            return render_template("single_fault.html")
        except BaseException as e:
            print(e)
            flash("密码错误","err")
        connection.commit()
        connection.close()

    return render_template("index.html")

#注册
@app.route("/register", methods=['GET', 'POST'])
def register():
    if request.method=='POST':
        uname=request.form.get("uname")
        upwd=request.form.get("upwd")
        print("uname",uname)
        print("upwd",upwd)
        connection = pymysql.connect(host='localhost',
                                     user='root',
                                     passwd='zhangyixin',
                                     db='faults_last',
                                     port=3306,
                                     charset='utf8'
                                     )
        cur = connection.cursor()  # 游标（指针）cursor的方式操作数据
        sql1 = 'SELECT * FROM users where username="' + uname + '" '
        try:
            cur.execute(sql1)  # execute(query, args):执行单条sql语句。
            see1 = cur.fetchone()  # 使结果全部可看
            print(see1[0])
            flash("用户名已存在", "error")
            return render_template("register.html")
        except BaseException as e:
            print(e)
            sql2 = 'insert into users(username,password) values("'+uname+'","'+upwd+'")'
            cur.execute(sql2)  # execute(query, args):执行单条sql语句。
            see1 = cur.fetchone()  # 使结果全部可看
            flash("注册成功", "ok")
        connection.commit()
        connection.close()
    return render_template("register.html")

@app.route("/test", methods=['GET', 'POST'])
def test():
    temp_init()
    return render_template("test.html")



#单故障录入页面
@app.route("/single_fault",methods=['GET', 'POST'])
def single():
    temp_init()
    return render_template("single_fault.html")

#单解决方案录入页面
@app.route("/single_solution",methods=['GET', 'POST'])
def single_solution():
    temp_init()
    return render_template("single_solution.html")

#流程图录入页面
@app.route("/flow_input",methods=['GET', 'POST'])
def flow_input():
    return render_template("flow.html")


#故障特征智能提取
@app.route("/zhin",methods=['GET', 'POST'])
def zhin():
    s=request.form['maintenance_plan']
    print('s_p', s)
    two=[]
    pattern = re.compile('发射车|装填车|运输车|装运车|FS车')
    car = pattern.search(s)
    if car is not None:
        print('car',car.group())
        s = s.replace(car.group(), '')
        two.append(car.group())
    else:
        two.append('')
    pattern1 = re.compile('1号车|2号车|3号车|4号车|5号车|6号车|7号车|8号车|9号车|10号车')
    type = pattern1.search(s)
    if type is not None:
        print('type',type.group())
        s=s.replace(type.group(),'')
        two.append(type.group())
    else:
        two.append('')
    print('two',two)
    weather=[]
    temperature = re.findall(r'-?\d+\.?\d*e?-?\d*?℃', s)
    if len(temperature)==0:
        temperature=['','']
    if len(temperature)==1:
        temperature.append('')
    humidity = re.findall(r'-?\d+\.?\d*e?-?\d*%', s)
    if len(humidity)==0:
        humidity.append('')
    pressure = re.findall(r'-?\d+\.?\d*e?-?\d*hpa', s)
    if len(pressure)==0:
        pressure.append('')
    height = re.findall(r'-?\d+\.?\d*e?-?\d*m', s)
    if len(height)==0:
        height.append('')
    salinity = re.findall(r'-?\d+\.?\d*e?-?\d*ppm', s)
    if len(salinity)==0:
        salinity.append('')
    # salinity=salinity.replace('ppm','')
    print('temperature', temperature)
    print('humidity', humidity)
    print('pressure', pressure)
    print('height', height)
    print('salinity', salinity)

    if temperature[0]:
        print('temperature[0]',temperature[0])
        s=s.replace(temperature[0],'')
        temperature[0]=temperature[0].replace('℃','')
    if temperature[1]:
        s=s.replace(temperature[1],'')
        temperature[1]=temperature[1].replace('℃','')
    if humidity[0]:
        s=s.replace(humidity[0],'')
        humidity[0]=humidity[0].replace('%','')
    if pressure[0]:
        s=s.replace(pressure[0],'')
        pressure[0]=pressure[0].replace('hpa','')
    if height[0]:
        s=s.replace(height[0],'')
        height[0]=height[0].replace('m','')
    if salinity[0]:
        s=s.replace(salinity[0],'')
        salinity[0]=salinity[0].replace('ppm','')

    print('s_b', s)

    weather.append(temperature)
    weather.append(humidity)
    weather.append(pressure)
    weather.append(height)
    weather.append(salinity)
    print('weather',weather)


    list_five_num=0
    if s =='':
        list_five=[]
    else:
        list_five=tfi(s)
        print('list_five',list_five)
        list_five_num=len(list_five)
        if list_five_num>0:
            print("识别成功")
    return json.dumps({'list_five': list_five,'list_five_num':list_five_num,'weather':weather,'two':two})


#故障特征录入
@app.route("/input_fault" ,methods=['GET', 'POST'])
def input_fault():
    global last_fea_count, last_sol_count,next_id
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 passwd='zhangyixin',
                                 db='faults_last',
                                 port=3306,
                                 charset='utf8'
                                 )
    cur = connection.cursor()  # 游标（指针）cursor的方式操作数据
    sql1 = 'SELECT count(*) FROM(SELECT count(*) FROM features GROUP BY features.group_id)a'
    cur.execute(sql1)  # execute(query, args):执行单条sql语句。
    see1 = cur.fetchone()  # 使结果全部可看
    print(see1[0])
    connection.commit()
    connection.close()
    next_id=see1[0]+1
    last_fea_count = Feature.query.filter(Feature.features_id).count()
    last_sol_count = Sol_feature.query.filter(Sol_feature.sol_features_id).count()
    print(last_fea_count)
    print(last_sol_count)
    feature1 = request.form.get("Program_features1")
    print('feature1',feature1)
    feature2 = request.form.get("Program_features2")
    feature3 = request.form.get("Program_features3")
    feature4 = request.form.get("Program_features4")
    feature5 = request.form.get("Program_features5")
    feature6 = request.form.get("Program_features6")
    feature7 = request.form.get("Program_features7")
    feature8 = request.form.get("Program_features8")
    feature9 = request.form.get("Program_features9")
    feature10 = request.form.get("Program_features10")

    weather1 = request.form.get("weather1")
    weather2 = request.form.get("weather2")
    weather3 = request.form.get("weather3")
    weather4 = request.form.get("weather4")
    weather5 = request.form.get("weather5")
    weather6 = request.form.get("weather6")
    weather7 = request.form.get("weather7")
    weather8 = request.form.get("weather8")

    new_weather1 = Feature(features=weather1, group_id=next_id)
    new_weather2 = Feature(features=weather2, group_id=next_id)
    new_weather3 = Feature(features=weather3, group_id=next_id)
    new_weather4 = Feature(features=weather4, group_id=next_id)
    new_weather5 = Feature(features=weather5, group_id=next_id)
    new_weather6 = Feature(features=weather6, group_id=next_id)
    new_weather7 = Feature(features=weather7, group_id=next_id)
    new_weather8 = Feature(features=weather8, group_id=next_id)

    db.session.add(new_weather1)
    db.session.add(new_weather2)
    db.session.add(new_weather3)
    db.session.add(new_weather4)
    db.session.add(new_weather5)
    db.session.add(new_weather6)
    db.session.add(new_weather7)
    db.session.add(new_weather8)
    db.session.commit()


    new_feature1 = Feature(features=feature1,group_id=next_id)



    f=open("KB_query\\external_dict\\features.txt",'a',encoding='utf-8')
    f.write(feature1+' features\n')


    db.session.add(new_feature1)
    try:
        if feature2!=None:
            new_feature2 = Feature(features=feature2,group_id=next_id)
            db.session.add(new_feature2)
            f.write(feature2 + ' features\n')
        if feature3!=None:
            new_feature3 = Feature(features=feature3,group_id=next_id)
            db.session.add(new_feature3)
            f.write(feature3 + ' features\n')
        if feature4!=None:
            new_feature4 = Feature(features=feature4,group_id=next_id)
            db.session.add(new_feature4)
            f.write(feature4 + ' features\n')
        if feature5!=None:
            new_feature5 = Feature(features=feature5,group_id=next_id)
            db.session.add(new_feature5)
            f.write(feature5 + ' features\n')
        if feature6!=None:
            new_feature6 = Feature(features=feature6,group_id=next_id)
            db.session.add(new_feature6)
            f.write(feature6 + ' features\n')
        if feature7!=None:
            new_feature7 = Feature(features=feature7,group_id=next_id)
            db.session.add(new_feature7)
            f.write(feature7 + ' features\n')
        if feature8!=None:
            new_feature8 = Feature(features=feature8,group_id=next_id)
            db.session.add(new_feature8)
            f.write(feature8 + ' features\n')
        if feature9!=None:
            new_feature9 = Feature(features=feature9,group_id=next_id)
            db.session.add(new_feature9)
            f.write(feature9 + ' features\n')
        if feature10!=None:
            new_feature10 = Feature(features=feature10,group_id=next_id)
            db.session.add(new_feature10)
            f.write(feature10 + ' features\n')
        f.close()
        db.session.commit()
    except BaseException as e:
        print("录入故障特征出错")
        print(e)
        # flash("录入出错","input_err")
    else:
        print("录入成功")
    return render_template("single_fault.html")

#解决方案智能识别
# @app.route("/zhin_sol",methods=['GET', 'POST'])
# def zhin_sol():
#     s=request.form.get("maintenance_plan")
#     print("EEEEEEEEEEEEEEEEe")
#     print('s',s)
#     list_five_num=0
#     if s =='':
#         list_five=[]
#     else:
#         list_five=tfi1(s)
#         print("识别tfi1")
#         print("QQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQqqq")
#         print(list_five)
#         list_five_num = len(list_five)
#         if list_five_num>0:
#             print("识别成功")
#             # flash("识别成功")
#     return json.dumps({'list_five': list_five, 'list_five_num': list_five_num})

#解决方案录入
@app.route("/input_sol" ,methods=['GET', 'POST'])
def input_sol():
    global next_id
    feature1 = request.form.get("maintenance_plan")

    new_sol1=Sol_feature(sol_features=feature1,group_id=next_id)
    db.session.add(new_sol1)
    f=open("KB_query\\external_dict\\solutions.txt",'a',encoding='utf-8')
    if feature1!=None:
        f.write(feature1+' solutions\n')

    db.session.commit()
    f.close()
    countf = Feature.query.filter(Feature.features_id).count()
    counts = Sol_feature.query.filter(Sol_feature.sol_features_id).count()

    global last_fea_count,last_sol_count
    print('last_fea_count',last_fea_count)
    print('last_sol_count',last_sol_count)
    try:
        for i in range(last_fea_count+1, countf+1):
            # for j in range(last_sol_count+1, counts+1):
            new_fs = Fea2Sol(features_id=i, sol_features_id=last_sol_count+1)
            db.session.add(new_fs)
            db.session.commit()
    except BaseException as e:
        print(e)


    print("提交成功")
#    create_nt(feature1)
    print("creat_nt")
#    upload_files(feature1)
    print("upload")
    flash("录入成功")
    return render_template("single_solution.html")

#批量录入页面
@app.route("/batch_fault",methods=['GET', 'POST'])
def batch_fault():
    phenomenon = []
    phenomenon_fenci = []
    reason = []
    reason_fenci = []
    method = []
    method_fenci = []

    global phenomenon_six,method_six
    phenomenon_six = []
    method_six = []
    global case_num
    case_num=0
    tb=[]

    if request.method=='POST':
        file = request.files['file']
        dst = os.path.join(os.path.dirname(__file__), file.filename)
        print('dst',dst)
        print(file.filename)
        name=file.filename
        suffix=name.split(".")
        print('suffix',suffix[1])
        if suffix[1]=='xlsx':
            book = xlrd.open_workbook(dst)
            # 获取所有的esheet
            list = book.sheets()
            sheet_names = book.sheet_names()
            print('sheet_names:',sheet_names)
            sheet_i = book.sheet_by_name('Sheet1')
            print(sheet_i)
            for i in sheet_names:
                sheet_i = book.sheet_by_name(i)
                print('sheet_i.ncols',sheet_i.ncols)
                if sheet_i.ncols==5:
                    case_num=sheet_i.nrows-1
                    # 创建一个for循环迭代读取xls文件每行数据的, 从第二行开始是要跳过标题行
                    for r in range(1, sheet_i.nrows):
                        phenomenon.append(sheet_i.cell(r, 1).value)
                        reason.append(sheet_i.cell(r, 2).value)
                        method.append(sheet_i.cell(r, 3).value)
                        # print('phenomenon['+str(r-1)+']',phenomenon[r-1])
                        # print('reason['+str(r-1)+']',reason[r-1])
                        # print('method['+str(r-1)+']',method[r-1])
                        temp = sheet_i.cell(r, 4).value + sheet_i.cell(r, 1).value
                        all = tiqu(temp)
                        if  phenomenon[r-1]== '':
                            phenomenon_fenci.append(phenomenon_six)
                        else:
                            phenomenon_six.append(all)
                            t2=tfi(phenomenon[r - 1])
                            for i in range(0,len(t2)):
                                phenomenon_six[r-1] .append(t2[i])
                            print('phenomenon_six', phenomenon_six)
                            phenomenon_fenci.append(phenomenon_six)
                        if method[r-1] == '':
                            method_fenci.append(method_six)
                        else:
                            method_six .append(method[r-1])
                            print('method_six',method_six)
                            method_fenci.append(method_six)


                elif sheet_i.ncols==7:
                    case_num = sheet_i.nrows-1
                    for r in range(1, sheet_i.nrows):
                        phenomenon.append(sheet_i.cell(r, 3).value)
                        reason.append(sheet_i.cell(r, 4).value)
                        method.append(sheet_i.cell(r, 5).value)
                        # print('phenomenon['+str(r-1)+']',phenomenon[r-1])
                        # print('reason['+str(r-1)+']',reason[r-1])
                        # print('method['+str(r-1)+']',method[r-1])
                        temp = sheet_i.cell(r, 1).value + sheet_i.cell(r, 6).value
                        all = tiqu(temp)
                        if  phenomenon[r-1]== '':
                            phenomenon_fenci.append(phenomenon_six)
                        else:
                            phenomenon_six.append(all)
                            t2 = tfi(phenomenon[r - 1])
                            for i in range(0, len(t2)):
                                phenomenon_six[r - 1].append(t2[i])
                            print('phenomenon_six', phenomenon_six)
                            phenomenon_fenci.append(phenomenon_six)
                        if method[r-1] == '':
                            method_fenci.append(method_six)
                        else:
                            method_six.append(method[r-1])
                            print('method_six',method_six)
                            method_fenci.append(method_six)
                elif sheet_i.ncols==15:
                    case_num = sheet_i.nrows-1
                    for r in range(1, sheet_i.nrows):
                        phenomenon.append(sheet_i.cell(r, 8).value)
                        phenomenon.append(sheet_i.cell(r, 9).value)
                        reason.append(sheet_i.cell(r, 10).value)
                        method.append(sheet_i.cell(r, 11).value)
                        # print('phenomenon['+str(r-1)+']',phenomenon[r-1])
                        # print('reason['+str(r-1)+']',reason[r-1])
                        # print('method['+str(r-1)+']',method[r-1])
                        temp = str(sheet_i.cell(r, 2).value) + str(sheet_i.cell(r, 3).value)+str(sheet_i.cell(r, 6).value) + str(sheet_i.cell(r, 7).value)+str(sheet_i.cell(r, 14).value)
                        all = tiqu(temp)
                        if  phenomenon[r-1]== '':
                            phenomenon_fenci.append(phenomenon_six)
                        else:
                            phenomenon_six.append(all)
                            t2 = tfi(phenomenon[r - 1])
                            for i in range(0, len(t2)):
                                phenomenon_six[r - 1].append(t2[i])
                            print('phenomenon_six', phenomenon_six)
                            phenomenon_fenci.append(phenomenon_six)

                        if method[r-1] == '':
                            method_fenci.append(method_six)
                        else:
                            method_six.append(method[r-1])
                            print('method_six',method_six)
                            method_fenci.append(method_six)
                num = len(phenomenon_fenci)
                print('num:',num)
                print('phenomenon_fenci',phenomenon_fenci)
                print('method_fenci',method_fenci)
        elif suffix[1]=='docx':
            file = request.files['file']
            dst = os.path.join(os.path.dirname(__file__), file.filename)
            doc = Document(dst)  # filename为word文档
            case_num = len(doc.tables)
            for i in range(0,case_num):
                tb.append("")
            print('tb',tb)
            tb[0] = doc.tables[0]
            rows_len = len(doc.tables[0].rows)
            print('rows_len', rows_len)
            cols_len = len(doc.tables[0].rows[0].cells)
            print('cols_len', cols_len)
            if cols_len==6:

                for i in range(1,case_num):
                    tb[i]=doc.tables[i]
                for i in range(0,case_num):
                    need1  = tb[i].rows[3].cells[0].text
                    print('need1', need1)
                    need2  = tb[i].rows[4].cells[0].text
                    print('need2', need2)
                    pheno=need1+need2
                    need3 = tb[0].rows[18].cells[0].text
                    print('need3', need3)
                    phenomenon_six.append(tfi(pheno))
                    print('phenomenon_six', phenomenon_six)
                    method_six .append(need3)
                    print('method_six', method_six)
            elif cols_len==4:
                case_num = len(doc.tables)
                for i in range(1,case_num):
                    tb[i]=doc.tables[i]
                for i in range(0,case_num):
                    pheno  = tb[0].rows[3].cells[0].text
                    print('pheno', pheno)
                    method = tb[0].rows[4].cells[0].text
                    print('method', method)
                    phenomenon_six .append(tfi(pheno))
                    print('phenomenon_six', phenomenon_six)
                    method_six .append(method)
                    print('method_six',method_six)
            elif cols_len==7:
                case_num = len(doc.tables)
                for i in range(1,case_num):
                    tb[i]=doc.tables[i]
                for i in range(0,case_num):
                    need1 = tb[0].rows[2].cells[4].text
                    need2 = tb[0].rows[2].cells[5].text
                    pheno=need1+need2
                    print('pheno',pheno)
                    method = tb[0].rows[2].cells[6].text
                    print('method', method)
                    phenomenon_six .append(tfi(pheno))
                    print('phenomenon_six', phenomenon_six)
                    method_six .append(method)
                    print('method_six', method_six)
    # create_nt()
    # upload_files()
    res = {}
    res['pheno'] = phenomenon_six
    res['method'] = method_six
    res['c'] = case_num
    return render_template("batch_fault.html")

#批量智能识别
@app.route("/batch_input",methods=['GET', 'POST'])
def batch_input():
    global phenomenon_six,method_six,case_num
    res = {}
    res['pheno'] = phenomenon_six
    res['method'] = method_six
    res['c'] = case_num
    print('res',res)
    # return json.dumps({'pheno':phenomenon_six,'method':method_six,'c':case_num})
    return res


#批量录入
@app.route("/batch_save",methods=['GET', 'POST'])
def batch_save():
    global case_num
    result_fea=[]
    result_sol=[]
    for i in range(0,case_num*2,2):
        val=request.values.get("ID"+str(i))
        val=val.split(",")
        result_fea.append(val)
    for i in range(1,case_num*2,2):
        val = request.values.get("ID" + str(i))
        result_sol.append(val)

    print('result_fea',result_fea)
    print('result_sol',result_sol)
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 passwd='zhangyixin',
                                 db='faults_last',
                                 port=3306,
                                 charset='utf8'
                                 )
    cur = connection.cursor()  # 游标（指针）cursor的方式操作数据
    sql1 = 'SELECT count(*) FROM(SELECT count(*) FROM features GROUP BY features.group_id)a'
    sql2 = 'SELECT count(*) FROM(SELECT count(*) FROM sol_features GROUP BY sol_features.group_id)b'
    cur.execute(sql1)
    see1 = cur.fetchone()
    print(see1[0])
    cur.execute(sql2)
    see2 = cur.fetchone()
    print(see2[0])
    connection.commit()
    connection.close()
    next_id = see1[0] + 1
    next_id2 = see2[0] + 1
    last_fea_count = Feature.query.filter(Feature.features_id).count()
    last_sol_count = Sol_feature.query.filter(Sol_feature.sol_features_id).count()
    print(last_fea_count)
    print(last_sol_count)
    for i in range(0,len(result_fea)):
        for j in range(0,len(result_fea[i])):
            new_feature=Feature(features=result_fea[i][j], group_id=next_id+i)
            db.session.add(new_feature)
            db.session.commit()

    for i in range(0,len(result_sol)):
        new_feature=Sol_feature(sol_features=result_sol[i], group_id=next_id2+i)
        db.session.add(new_feature)
        db.session.commit()


    countf = Feature.query.filter(Feature.features_id).count()
    counts = Sol_feature.query.filter(Sol_feature.sol_features_id).count()


    for i in range(last_fea_count+1, countf+1):
        for j in range(last_sol_count+1, counts+1):
            new_fs = Fea2Sol(features_id=i, sol_features_id=j)
            db.session.add(new_fs)
            db.session.commit()

    return render_template("batch_fault.html")

#流程图录入
@app.route("/flowchart", methods=['GET', 'POST'])
def flowchart():
    global list_chart, links
    list_chart = []
    links = []
    global list_chart1, links1
    list_chart1 = []
    links1 = []
    if request.method == 'POST':
        file = request.files['file']
        dst = os.path.join(os.path.dirname(__file__), file.filename)
        print('dst', dst)
        print('name', file.filename)


        image = Image.open(dst)
        image_list = cut_image(image)
        save_images(image_list)

        img_net = 'output2.png'
        img_net = Image.open(img_net).convert('RGB')
        # img.show()
        img_net = np.array(img_net)
        text = text_predict(img_net)
        print('text', text)

        img_net1 = 'output1.png'
        img_net1 = Image.open(img_net1).convert('RGB')
        # img.show()
        img_net1 = np.array(img_net1)
        text1 = text_predict(img_net1)

        print('text1[text]', text1)
        text1_list=[]
        for tl in range(0,len(text1)):
            text1_list.append(text1[tl]['text'])
        print('text1_list',text1_list)
        text1_one=tiqu(text1_list[0])
        text1_two=tfi(text1_list[0])
        for tt in range(0,len(text1_two)):
            text1_one.append(text1_two[tt])
        print('text1_two',text1_one)

        cx_list = []
        cy_list = []
        w_list = []
        h_list = []
        text_list = []
        for i in range(0, len(text)):
            cx_list.append(int(text[i]['cx']))
            cy_list.append(int(text[i]['cy']))
            w_list.append(int(text[i]['w']))
            h_list.append(int(text[i]['h']))
            text_list.append(text[i]['text'])
            print("textcx", int(text[i]['cx']))
            print("textcy", int(text[i]['cy']))
            print('textsub: ', text[i]['text'])
        flag_del = []
        for m in range(0, len(cy_list) - 1):
            for n in range(m + 1, len(cy_list)):
                if (cy_list[n] - cy_list[m] < 40 and abs(cx_list[n] - cx_list[m]) < 20) | (
                                    cy_list[n] - cy_list[m] < 2 and abs(cx_list[n] - cx_list[m]) < 80):
                    flag = 0
                    for k in range(0, len(flag_del)):
                        if (m in flag_del[k] and n not in flag_del[k]):
                            flag = 1
                            flag_del[k].append(n)
                        elif (m not in flag_del[k] and n in flag_del[k]):
                            flag = 1
                            flag_del[k].append(m)
                        elif (m in flag_del[k] and n in flag_del[k]):
                            flag = 1
                    if flag == 0:
                        flag_del.append([m, n])

        print('flag_del', flag_del)
        flag_del_temp = []
        for i in range(0, len(flag_del)):
            for j in range(1, len(flag_del[i])):
                flag_del_temp.append(flag_del[i][j])
        print('flag_del_temp', flag_del_temp)
        print('cx', cx_list)
        print('cy', cy_list)
        character_xy = []
        for i in range(0, len(cx_list)):
            if i not in flag_del_temp:
                character_xy.append([cx_list[i], cy_list[i]])

        print('text', text_list)
        print('character_xy', character_xy)  # 合并后的文字坐标数组
        text_concat = []
        for k in range(0, len(text_list)):
            flag = 0
            for i in range(0, len(flag_del)):
                if k in flag_del[i]:
                    flag = 1
                if k == flag_del[i][0]:
                    str = ""
                    for j in range(0, len(flag_del[i])):
                        str += text_list[flag_del[i][j]]
                    text_concat.append(str)
            if flag == 0:
                text_concat.append(text_list[k])

        print("text_concat: ", text_concat)  # 合并后的文字数组
        print('w', w_list)
        print('h', h_list)
        print('文本预测:', list(map(lambda x: x['text'], text)))

        # 火车票
        res_net = trainTicket.trainTicket(text)
        print('火车票预测:', res_net)

        # 身份证
        res_net = idcard.idcard(text)
        print('身份证预测:', res_net)

        for i in range(0, len(cx_list)):
            list_chart1.append({'name': i + 1, 'value': 0, 'x': cx_list[i], 'y': cy_list[i]})

        img = cv2.imread('output2.png')
        # 灰度
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # 二值化
        ret, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY)
        # 取反
        bitwiseNot = cv2.bitwise_not(binary)
        cv2.imshow("Not", bitwiseNot)
        cv2.imwrite('Not.png', bitwiseNot)
        contours, hierarchy = cv2.findContours(bitwiseNot, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        len_contours = len(contours)
        cv2.drawContours(img, contours[len_contours - 1], -1, (0, 0, 0), 3)
        # cv2.drawContours(img,contours,0,(0,255,0),3)
        print(type(contours))
        print(type(contours[0]))
        print(len(contours))
        cv2.fillPoly(img, contours, (0, 0, 0))

        bitwiseNot2 = cv2.bitwise_not(img)
        cv2.imshow("Not2", bitwiseNot2)
        cv2.imwrite('Not2.png', bitwiseNot2)
        # result=open_demo(img)
        kernel = np.ones((14, 14), np.uint8)
        kernel1 = np.ones((16, 16), np.uint8)
        result = cv2.erode(bitwiseNot2, kernel)
        cv2.imshow("img", result)
        rect = cv2.dilate(result, kernel1)
        cv2.imshow("rect", rect)
        cv2.imwrite('rect.png', rect)
        arrow = cv2.subtract(bitwiseNot2, rect, dst=None, mask=None, dtype=None)
        cv2.imshow("arrow", arrow)
        cv2.imwrite('arrow.png', arrow)
        im_gray = cv2.cvtColor(arrow, cv2.COLOR_RGB2GRAY)
        ret, thresh_arrow = cv2.threshold(im_gray, 128, 255, cv2.THRESH_BINARY)
        cv2.imwrite('thresh_arrow.png', thresh_arrow)

        nccomps = cv2.connectedComponentsWithStats(thresh_arrow)  # labels,stats,centroids
        print('nccomps', nccomps)
        _ = nccomps[0]
        print('_', _)
        labels = nccomps[1]
        print('labels', labels)
        centroids = nccomps[3]
        print('centroids', centroids)
        status = nccomps[2]
        print('status', status)
        for row in range(status.shape[0]):
            if status[row, :][0] == 0 and status[row, :][1] == 0:
                background = row
            else:
                continue
        status_no_background = np.delete(status, background, axis=0)
        rec_value_max = np.asarray(status_no_background[:, 4].max())
        print('rec_value_max', rec_value_max)
        # res=np.asarray(status_no_background[:,2:4])
        res = []
        res_index = []
        for i in range(0, len(status_no_background)):
            if abs(status_no_background[i][3] - status_no_background[i][2]) > 5 and status_no_background[i][3] > 9:
                # if abs(status_no_background[i][3] - status_no_background[i][2]) > 5:
                res.append(status_no_background[i])
                res_index.append(i)
        print('res', res)
        print('res_index', res_index)
        res_arr = np.asarray(res, 'int32')
        print(len(res))
        print('res_arr', res_arr)
        delete_res_arr = []
        delete_res_arr_index = []
        for m in range(0, len(res_arr)):
            if (res_arr[m][2] > 350):
                delete_res_arr.append(res_arr[m])
                delete_res_arr_index.append(m)
                np.delete(res_arr, m, axis=0)
        print('delete_res_arr', delete_res_arr)
        print('delete_res_arr_index', delete_res_arr_index)

        # re_value_max_position = np.asarray(status_no_background[:, 4].argmax())
        # print('re_value_max_position', re_value_max_position)
        h = np.asarray(labels, 'uint8')
        print('h', len(h))
        print('h[0]', len(h[0]))
        for i in range(0, len(res)):
            h[h == (res_index[i] + 1)] = 255
        kernel_h = np.ones((3, 3), np.uint8)
        result_h = cv2.dilate(h, kernel_h)
        cv2.imshow('h', h)
        cv2.imwrite('h.png', h)
        cv2.imshow('result_h', result_h)
        cv2.imwrite('result_h.png', result_h)
        # cv2.waitKey(0)

        # 设置并解析参数
        ap = argparse.ArgumentParser()
        ap.add_argument("-i", "--image", default='rect.png', help="path to the input image")
        args = vars(ap.parse_args())

        # 读取图片
        image = cv2.imread(args["image"])
        # 进行裁剪操作
        resized = imutils.resize(image, width=300)
        ratio = image.shape[0] / float(resized.shape[0])
        # 进行高斯模糊操作
        blurred = cv2.GaussianBlur(resized, (5, 5), 0)
        # 进行图片灰度化
        gray = cv2.cvtColor(blurred, cv2.COLOR_BGR2GRAY)
        # 进行颜色空间的变换
        lab = cv2.cvtColor(blurred, cv2.COLOR_BGR2LAB)
        # 进行阈值分割
        thresh = cv2.threshold(gray, 60, 255, cv2.THRESH_BINARY)[1]
        # 在二值图片中寻找轮廓
        cnts = cv2.findContours(thresh.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        cnts = imutils.grab_contours(cnts)

        chart_x = []
        chart_y = []
        chart_xy = []
        try:
            # 遍历每一个轮廓
            for c in cnts:
                # 计算每一个轮廓的中心点
                M = cv2.moments(c)
                # print('M',M)
                cX = int((M["m10"] / M["m00"]) * ratio)
                # print('cX', cX)
                chart_x.append(cX)
                cY = int((M["m01"] / M["m00"]) * ratio)
                # print('cY',cY)
                chart_y.append(cY)
                chart_xy.append([cX, cY])
                # 进行坐标变换
                c = c.astype("float")
                c *= ratio
                c = c.astype("int")
                # 绘制轮廓并显示结果
                cv2.drawContours(image, [c], -1, (0, 255, 0), 2)
                # cv2.putText(image, text, (cX, cY), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 255, 255), 2)
                cv2.imshow("Image", image)
            # cv2.waitKey(0)
            # print('chart_x',chart_x)
            # print('chart_y',chart_y)
            print('chart_xy', chart_xy)
            # chart_xy [[485, 570], [65, 570], [255, 525], [255, 393], [688, 393], [499, 341], [499, 235], [91, 228], [311, 129], [311, 27]]
            # line_xy [[448, 530], [103, 530], [254, 474], [254, 361], [664, 352], [383, 340], [498, 291], [498, 202], [423, 128], [136, 154], [310, 78]]
            # distance_min [[2, 0], [2, 1], [3,  12], [7, 3], [5, 4], [6,  5], [6,  5], [8, 6], [9, 8], [8, 7], [9, 8]]
            # 将文字填入框中
            ccdist = {}  # 保存文字和框对应关系的列表
            # visit = []  # 标记有没有被访问过
            for i in range(len(character_xy)):
                min = 0                 #保存最小下标
                dist = sys.maxsize      #保存最小距离
                for j in range(len(chart_xy)):
                    if character_xy[i][1] > chart_xy[j][1] : continue
                    tdist = abs((character_xy[i][0] - chart_xy[j][0]) * (character_xy[i][0] - chart_xy[j][0]) + (
                    character_xy[i][1] - chart_xy[j][1]) * (character_xy[i][1] - chart_xy[j][1]))
                    if tdist < dist:
                        min = j
                        dist = tdist
                # print('min:i ',min,i)
                # visit.append(min)
                ccdist[min] = i
            print('ccdist: ', ccdist)   #框和文字的映射关系字典

            chart_xy_reverse = []
            for i in range(0, len(chart_xy)):
                chart_xy_reverse.append(chart_xy[len(chart_xy) - i - 1])
            print('chart_xy_reverse', chart_xy_reverse)
            # cv2.waitKey()

            # 设置并解析参数
            ap = argparse.ArgumentParser()
            ap.add_argument("-i", "--image", default='./result_h.png', help="path to the input image")
            args = vars(ap.parse_args())

            # 读取图片
            image = cv2.imread(args["image"])
            # 进行裁剪操作
            resized = imutils.resize(image, width=300)
            ratio = image.shape[0] / float(resized.shape[0])

            # 进行高斯模糊操作
            blurred = cv2.GaussianBlur(resized, (5, 5), 0)
            # 进行图片灰度化
            gray = cv2.cvtColor(blurred, cv2.COLOR_BGR2GRAY)
            # 进行颜色空间的变换
            lab = cv2.cvtColor(blurred, cv2.COLOR_BGR2LAB)
            # 进行阈值分割
            thresh = cv2.threshold(gray, 60, 255, cv2.THRESH_BINARY)[1]
            # cv2.imshow("Thresh", thresh)
            # 在二值图片中寻找轮廓
            cnts = cv2.findContours(thresh.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            # print("**************************************")
            # print('cnts',cnts)
            cnts = imutils.grab_contours(cnts)
            # print('cnts',cnts)
            line_xy = []
            # 遍历每一个轮廓
            len_change = []
            temp_count = -1
            for c in cnts:
                temp_count = temp_count + 1
                # 计算每一个轮廓的中心点
                M = cv2.moments(c)
                L = cv2.arcLength(c, False)
                if L > 300:
                    len_change.append(temp_count)

                # print('M',M)
                cX = int((M["m10"] / M["m00"]) * ratio)
                # print('cX',cX)
                cY = int((M["m01"] / M["m00"]) * ratio)
                # print('cY',cY)
                line_xy.append([cX, cY])
                # 进行坐标变换
                c = c.astype("float")
                c *= ratio
                c = c.astype("int")
                cv2.drawContours(image, [c], -1, (0, 255, 0), 2)
                cv2.imshow("Image", image)
            # cv2.waitKey()
            print('len_change', len_change)
            print('line_xy', line_xy)
        except BaseException as e:
            print("error")
            print(e)
            flash("无法识别", "err")
        else:
            # 线段中心点距矩形中心点的距离
            distance_res = []
            bigY = []
            smaY = []
            bigX = []
            smaX = []
            for i in range(0, len(line_xy)):
                temp = []
                tybig = []
                tysma = []
                txbig = []
                txsma = []
                for j in range(0, len(chart_xy)):
                    # 将y坐标分出大小存放
                    if chart_xy[j][1] > line_xy[i][1]:
                        tybig.append(j)
                    else:
                        tysma.append(j)

                    # 将x坐标分出大小存放
                    if chart_xy[j][0] > line_xy[i][0]:
                        txbig.append(j)
                    else:
                        txsma.append(j)

                    distance = pow(line_xy[i][0] - chart_xy[j][0], 2) + pow(line_xy[i][1] - chart_xy[j][1], 2)
                    temp.append(distance)
                distance_res.append(temp)
                bigY.append(tybig)
                smaY.append(tysma)
                bigX.append(txbig)
                smaX.append(txsma)
                print('temp: ', temp)
                print('tybig: ', tybig)
                print('tysma: ', tysma)
                print('txbig: ', txbig)
                print('txsma: ', txsma)
            # 每条线段距离每个矩形的距离
            print('distance_res', distance_res)
            # for i in range(0,len(delete_res_arr_index)):
            #     distance_res.pop(delete_res_arr_index[i])


            delete_res_distance = []
            # 每条超长线段的端点距每个矩形的距离
            for i in range(0, len(delete_res_arr)):
                min_distance1 = 100000000
                min_distance1_index = -1
                min_distance2 = 100000000
                min_distance2_index = -1
                for j in range(0, len(chart_xy)):
                    distance1 = pow(delete_res_arr[i][0] - chart_xy[j][0], 2) + pow(
                        delete_res_arr[i][1] - chart_xy[j][1],
                        2)
                    distance2 = pow(delete_res_arr[i][0] + delete_res_arr[i][2] - chart_xy[j][0], 2) + pow(
                        delete_res_arr[i][1] + delete_res_arr[i][3] - chart_xy[j][1],
                        2)
                    if (distance1 < min_distance1):
                        min_distance1 = distance1
                        min_distance1_index = j
                    if (distance2 < min_distance2):
                        min_distance2 = distance2
                        min_distance2_index = j
                delete_res_distance.append([min_distance2_index, min_distance1_index])
            print('delete_res_distance', delete_res_distance)

            distance_min = []
            # 找出每条线段距离最近的两个矩形
            for i in range(0, len(distance_res)):


                mini = 0
                maxi = 0
                bigdict = {}
                smadict = {}
                for j in range(len(bigY[i])):
                    bigdict[bigY[i][j]] = distance_res[i][bigY[i][j]]
                    # if distance_res[i][bigY[i][j]] < distance_res[i][bigY[i][maxi]] :
                    #     maxi = j
                for k in range(len(smaY[i])):
                    smadict[smaY[i][k]] = distance_res[i][smaY[i][k]]
                    # if distance_res[i][smaY[i][k]] < distance_res[i][smaY[i][mini]] :
                    #     mini = k
                sortbig = sorted(bigdict.items(), key=lambda item: item[1])
                sortsma = sorted(smadict.items(), key=lambda item: item[1])
                print('sortbig: ', sortbig)
                print('sortsma: ', sortsma)
                m = 0
                n = 0
                print('lenbig', len(sortbig))
                print('lensma', len(sortsma))

                print("m, n: ", m, n);
                if sortbig[m][0] > sortsma[n][0]:
                    distance_min.append([sortbig[m][0], sortsma[n][0]])
                else:
                    distance_min.append([sortsma[m][0], sortbig[n][0]])

            print('distance_min', distance_min)
            if delete_res_distance != []:
                distance_min.append(delete_res_distance[0])

            first_count, second_count = [], []
            for i in range(0, len(chart_xy)):
                first_count.append(0)
                second_count.append(0)
            for i in range(0, len(distance_min)):
                first_count[distance_min[i][0]] = first_count[distance_min[i][0]] + 1
                second_count[distance_min[i][1]] = second_count[distance_min[i][1]] + 1
            print('first_count', first_count)
            print('second_count', second_count)
            end = []
            for i in range(len(first_count)):
                if (first_count[i]) == 0:
                    end.append(i)
            global case_num
            case_num = first_count.count(0)
            for i in range(0, len(second_count) - 1):
                if second_count[i] == 0:
                    for j in range(0, len(distance_min)):
                        if distance_min[j][0] == i:
                            distance_min[j][0], distance_min[j][1] = distance_min[j][1], distance_min[j][0]
                    case_num = case_num + 1

            print('case_num', case_num)
            print('distance_min', distance_min)
            if len(chart_xy_reverse) == len(character_xy):
                print('=')
            elif len(chart_xy_reverse) > len(character_xy):
                for i in range(len(character_xy), len(chart_xy_reverse)):
                    character_xy.insert(i, [0, 0])
                    text_concat.insert(i, i)
                print('character_xy', character_xy)
            else:
                print('<')

            # 存放填补后的文字
            character_result = []

            for i in range(0, len(character_xy)):
                # name = "第1526行代码问题"
                # if i in ccdist:
                # name = text_concat[ccdist[i]]
                #     character_result.append(name)
                list_chart.append({'name': text_concat[len(character_xy)-i-1], 'value': 0, 'x': character_xy[len(character_xy)-i-1][0], 'y': character_xy[len(character_xy)-i-1][1], 'id': i}, )

            print('list_chart', list_chart)
            print('character_result', character_result)

            for m in range(0, len(distance_min)):
                links.append({'source': distance_min[m][0], 'target': distance_min[m][1], })

            g = Graph()  # 实例化图类
            try:
                for i in range(len(chart_xy)):
                    g.addVertex(i, text_concat[i])  # 给邻接表添加节点

                for v in g:
                    print('v.getText', v.getText())
                print(g.vertList)  # 打印邻接表

                for j in range(len(distance_min)):
                    g.addEdge(distance_min[j][0], distance_min[j][1], 0)  # 给邻接表添加边及权重

                res = []
                graph = {}
                for i in range(len(chart_xy)):
                    graph[i] = []
                print(graph)
                for v in g:  # 循环每个顶点
                    for w in v.getConnections():  # 循环每个顶点的所有邻居节点
                        #    print("(%s, %s)" % (v.getId(), w.getId())) #打印顶点和其邻居节点的键
                        res.append([v.getId(), w.getId()])

                print('res', res)

                for i in range(0, len(res)):
                    graph[res[i][0]].append(res[i][1])
                print('graph', graph)
                global allpath
                allpath = []
                for i in range(len(end)):
                    allpath.append(findPath(graph, len(chart_xy) - 1, end[i]))
                    reversed(allpath)

                print('\n所有路径：', allpath)

                print('ccdist', ccdist)
                allpath_new = []
                for i in range(len(allpath)):
                    temp_a = []
                    for j in range(len(allpath[i])):
                        if allpath[i][j] in ccdist.keys():
                            temp_a.append(ccdist[allpath[i][j]])
                    allpath_new.append(temp_a)
                print('allpath_new', allpath_new)
                allpath = allpath_new
                print('character_result', character_result)

                # 所有节点的文字
                global case_text
                case_text = []
                for i in range(0, len(allpath)):
                    temp = []
                    for j in range(0, len(allpath[i])):

                        for v in g:
                            if v.getId() == allpath[i][j]:
                                temp.append(v.getText())
                    case_text.append(temp)
                print('case_text', case_text)



            except BaseException as e:
                print("error")
                print(e)

    return render_template("flowchart.html")

#流程图数据

@app.route("/flowData", methods=['GET', 'POST'])
def flowData():
    global list_chart,links,case_num,allpath,case_text
    return json.dumps({'list_chart':list_chart,'links':links,'case_num':case_num,'allpath':allpath,'case_text':case_text})
    # return render_template("flowchart.html",list_chart=list_chart,links=links)



@app.route("/test",methods=['GET', 'POST'])
def test1():
    global fuzz1,fuzz_q1
    global f1,f2,f3,f4,f5,f6
    fuzz_features=""
    print(f1)
    print(f2)
    if f1 !='':
        fuzz_features+=f1+"&"
    if f2 !='':
        fuzz_features += f2 + "&"
    if f3 !='':
        fuzz_features += f3 + "&"
    if f4 !='':
        fuzz_features += f4 + "&"
    if f5 !='':
        fuzz_features += f5 + "&"
    if f6 !='':
        fuzz_features += f6 + "&"
    fuzz_len=len(fuzz1)

    list=[]
    list.append({"name": '特征:' + fuzz_features,"des": 'nodedes01', "symbolSize": 70, "category": 0, }, )
    for i in range(fuzz_len):
        list.append({"name": '案例' + str(i + 1),"des": 'nodedes01', "symbolSize": 50, "category": i+1,}, )
        list.append({"name": str(i + 1)+'特征' ,"des": 'nodedes01', "symbolSize": 50, "category":i+1, }, )
        list.append({"name":  str(i + 1)+'解决方案' ,"des": 'nodedes01', "symbolSize": 50, "category": i+1, }, )

    for j in range(fuzz_len):
        for k in range(len(fuzz1[j])):
            list.append({"name":  fuzz1[j][k], "des": 'nodedes01',"symbolSize": 50, "category": j+1,}, )
    for a in range(fuzz_len):
        for b in range(len(fuzz_q1[a])):
            list.append({"name": fuzz_q1[a][b],"des": 'nodedes01', "symbolSize": 50, "category": a+1,}, )

    links=[]
    for i in range(fuzz_len):
        links.append({"source": '特征:' + fuzz_features, "target": '案例' + str(i + 1),"name":' '}, )
        links.append({"source": '案例' + str(i + 1), "target": str(i + 1) + '特征',"name":' '}, )
        links.append({"source": '案例' + str(i + 1), "target": str(i + 1) + '解决方案',"name":' '}, )
    for j in range(fuzz_len):
        for k in range(len(fuzz1[j])):
            links.append({"source": str(j + 1) + '解决方案', "target": fuzz1[j][k], 'name': '特征' + str(k + 1)}, )
    for a in range(fuzz_len):
        for b in range(len(fuzz_q1[a])):
            links.append({"source": str(a + 1) + '特征', "target": fuzz_q1[a][b],"name":'特征'+str(b+1),} )

    print(list)
    print(links)
    return json.dumps({"list": list, "links": links})



#流程图录入(弃）
# @app.route("/flow_zhin",methods=['GET', 'POST'])
# def flow_zhin():
#     list1=["F计算机","1-4圆筒位置为空","1-4圆筒Q飞灯不亮"]
#     list2=["更换","F电源保险管",""]
#     list3=["F计算机","1-4圆筒位置为空","1-4圆筒Q飞灯不亮","更换F电源保险管","无效"]
#     list4=["更换","F电源控制模块",""]
#     list5=["F计算机","1-4圆筒位置为空","1-4圆筒Q飞灯不亮","更换F电源保险管","F电源控制模块","无效"]
#     list6=["更换","F电源模块",""]
#     return render_template("flow.html",a11=list1[0],a12=list1[1],a13=list1[2],
#                            b11=list2[0],b12=list2[1],b13=list2[2],
#                            a21=list3[0],a22=list3[1],a23=list3[2],a24=list3[3],a25=list3[4],
#                            b21=list4[0],b22=list4[1],b23=list4[2],
#                            a31=list5[0],a32=list5[1],a33=list5[2],a34=list5[3],a35=list5[4],a36=list5[5],
#                            b31=list6[0],b32=list6[1],b33=list6[2]
#                            )

#模糊查询页面

#精确/模糊查询特征智能提取
@app.route("/zhin_query",methods=['GET', 'POST'])
def zhin_query():
    s=request.form['maintenance_plan']
    print('s_p', s)
    two=[]
    pattern = re.compile('发射车|装填车|装运车|运输车')
    car = pattern.search(s)
    if car is not None:
        print('car',car.group())
        s = s.replace(car.group(), '')
        two.append(car.group())
    else:
        two.append('')
    pattern1 = re.compile('1号车|2号车|3号车|4号车|5号车|6号车|7号车|8号车|9号车|10号车')
    type = pattern1.search(s)
    if type is not None:
        print('type',type.group())
        s=s.replace(type.group(),'')
        two.append(type.group())
    else:
        two.append('')
    print('two',two)
    weather=[]
    temperature = re.findall(r'-?\d+\.?\d*e?-?\d*?℃', s)
    if len(temperature)==0:
        temperature=['','']
    if len(temperature)==1:
        temperature.append('')
    humidity = re.findall(r'-?\d+\.?\d*e?-?\d*%', s)
    if len(humidity)==0:
        humidity.append('')
    pressure = re.findall(r'-?\d+\.?\d*e?-?\d*hpa', s)
    if len(pressure)==0:
        pressure.append('')
    height = re.findall(r'-?\d+\.?\d*e?-?\d*m', s)
    if len(height)==0:
        height.append('')
    salinity = re.findall(r'-?\d+\.?\d*e?-?\d*ppm', s)
    if len(salinity)==0:
        salinity.append('')
    # salinity=salinity.replace('ppm','')
    print('temperature', temperature)
    print('humidity', humidity)
    print('pressure', pressure)
    print('height', height)
    print('salinity', salinity)
    if temperature[0]:
        print('temperature[0]',temperature[0])
        s=s.replace(temperature[0],'')
    if temperature[1]:
        s=s.replace(temperature[1],'')
    if humidity[0]:
        s=s.replace(humidity[0],'')
    if pressure[0]:
        s=s.replace(pressure[0],'')
    if height[0]:
        s=s.replace(height[0],'')
    if salinity[0]:
        s=s.replace(salinity[0],'')



    print('s_b', s)

    weather.append(temperature)
    weather.append(humidity)
    weather.append(pressure)
    weather.append(height)
    weather.append(salinity)
    print('weather',weather)


    list_five_num=0
    if s =='':
        list_five=[]
    else:
        list_five=tfi(s)
        print('list_five',list_five)
        list_five_num=len(list_five)
        if list_five_num>0:
            print("识别成功")
    return json.dumps({'list_five': list_five,'list_five_num':list_five_num,'weather':weather,'two':two})


@app.route("/fuzz_query",methods=['GET', 'POST'])
def fuzz():
    # global y1, y2, y3, y4, y5, y6, s1, s2, s3, s4, s5, s6,num1, num2, num3, num4, num5, num6
    # # 精确查询中的特征及解决办法
    # y1, y2, y3, y4, y5, y6 = " ", " ", " ", " ", " ", " "
    # s1, s2, s3, s4, s5, s6 = "", "", "", "", "", ""
    # # 精确查询中每个特征在数据库中的数量
    # num1, num2, num3, num4, num5, num6 = "", "", "", "", "", ""
    return render_template("fuzz_query.html")


#模糊查询
@app.route("/submit_fuzz",methods=['GET', 'POST'])
def submit_fuzz():
    if request.method == 'POST':
        global  v1
        v1 = request.form.get("v1")
        print("v1:",v1)
        fea_corpus = []
        sol_corpus = []
        f_fea = open('./dict/features.txt', encoding='utf-8')
        for line in f_fea:
            print(line.strip())
            fea_corpus.append(line.strip())
        f_fea.close()

        f_sol = open('./dict/solutions.txt', encoding='utf-8')
        for line in f_sol:
            print(line.strip())
            sol_corpus.append(line.strip())
        print(sol_corpus)
        f_sol.close()

        global f1, f2, f3, f4, f5,f6,f7,f8,f9,f10,w1,w2,w3,w4,w5,w6,w7,w8,max_G,max_G_id
        f1,f2,f3,f4,f5,f6,f7,f8,f9,f10=None,None,None,None,None,None,None,None,None,None
        w1, w2, w3, w4, w5, w6, w7, w8= '', None, None, None, None, None, None, None
        f1 = request.form.get("fault_features1");
        print('f1',f1)
        f2 = request.form.get("fault_features2")
        f3 = request.form.get("fault_features3")
        f4 = request.form.get("fault_features4")
        f5 = request.form.get("fault_features5")
        f6 = request.form.get("fault_features6")
        f7 = request.form.get("fault_features7")
        f8 = request.form.get("fault_features8")
        f9 = request.form.get("fault_features9")
        f10 = request.form.get("fault_features10")
        w1 = request.form.get("weather1")
        print('w1::',w1)
        w2 = request.form.get("weather2")
        w3 = request.form.get("weather3")
        w4 = request.form.get("weather4")
        w5 = request.form.get("weather5")
        w6 = request.form.get("weather6")
        w7 = request.form.get("weather7")
        w8 = request.form.get("weather8")
        # f1, f2, f3, f4, f5,f6=feature1,feature2,feature3,feature4,feature5,feature6
    #    features = feature1 + feature2 + feature3 + feature4 + feature5+feature6
    #    similar_list, ori_question_dict, ori_answer_dict = get_corpus_consine(features)  # 计算所有相似度
    #    result=get_three_answer(similar_list,ori_question_dict,ori_answer_dict)
        temp=' '
        if w1:
            print("%%%%%%%%%%%%%%%%%%%%")
            print('w1',w1)
            print('temp',temp)
            temp+=w1+' '
        if w2:
            temp += w2 + ' '
        if w3:
            temp += w3 + ' '
        if w4:
            temp += w4 + ' '
        if w5:
            temp += w5 + ' '
        if w6:
            temp += w6 + ' '
        if w7:
            temp += w7 + ' '
        if w8:
            temp += w8 + ' '
        if f1:
            temp+=f1+' '
        if f2:
            temp += f2 + ' '
        if f3:
            temp += f3 + ' '
        if f4:
            temp += f4 + ' '
        if f5:
            temp += f5 + ' '
        if f6:
            temp += f6 + ' '
        if f7:
            temp += f7 + ' '
        if f8:
            temp += f8 + ' '
        if f9:
            temp += f9 + ' '
        if f10:
            temp += f10 + ' '
        print('temp',temp)
        fea_corpus.append(temp)

        # 将文本中的词语转换为词频矩阵
        vectorizer = CountVectorizer()
        # 计算个词语出现的次数
        X = vectorizer.fit_transform(fea_corpus)  # 获取词袋中所有文本关键词
        word = vectorizer.get_feature_names()
        # 类调用
        transformer = TfidfTransformer()

        # 将词频矩阵X统计成TF-IDF值
        tfidf = transformer.fit_transform(X)
        # 查看数据结构 tfidf[i][j]表示i类文本中的tf-idf权重
        weight = tfidf.toarray()
        # print (weight)
        len_w = len(weight)
        # kmeans聚类
        kmeans = KMeans(n_clusters=6, random_state=0).fit(weight)  # k值可以自己设置
        centroid_list = kmeans.cluster_centers_
        pca = PCA(n_components=2)
        result = pca.fit_transform(centroid_list)
        result1 = pca.fit_transform(weight)
        print(weight[len_w - 1])

        labels = kmeans.labels_
        n_clusters_ = len(centroid_list)
        # print "cluster centroids:",centroid_list
        len_labels = len(labels)
        labels_last = labels[len_labels - 1]
        print('labels_last', labels_last)
        print(labels)
        max_centroid = 0
        max_cluster_id = 0
        cluster_menmbers_list = []
        for i in range(0, n_clusters_):
            menmbers_list = []
            for j in range(0, len(labels)):
                if labels[j] == i:
                    menmbers_list.append(j)
            cluster_menmbers_list.append(menmbers_list)
        # print cluster_menmbers_list
        # 聚类结果
        center_list = []
        print("cluster_members_list", cluster_menmbers_list)
        for i in range(0, len(cluster_menmbers_list)):
            print('第' + str(i) + '类' + '---------------------')
            if i == labels_last:
                center_list = result1[cluster_menmbers_list[i]]
                print(center_list)
                print('len(center_list)',len(center_list))
                print('result1[cluster_member_list[i]]', result1[cluster_menmbers_list[i]])
            for j in range(0, len(cluster_menmbers_list[i])):
                a = cluster_menmbers_list[i][j]
                print(fea_corpus[a])
        # print('center_list',center_list)
        x = result1[len_w - 1:len_w, :]

        print('x',x)
        rbf = RBF(2, len(center_list), 1, center_list)
        rbf.train(x)
        # z=rbf.predict(x)
        # print('z',z)

        print('cluster_menmbers_list', cluster_menmbers_list)
        print('max_G_id', max_G_id)

        dict_zip = dict(zip(cluster_menmbers_list[labels_last], similar_list))
        print('dict_zip', dict_zip)
        after_sort_dict = sorted(dict_zip.items(), key=lambda x: x[1], reverse=True)
        print('after_sort_dict', after_sort_dict)
        after_sort_dict_index.clear()
        for i in range(len(after_sort_dict)):
            after_sort_dict_index.append(after_sort_dict[i][0])

        print('after_sort_dict_index', after_sort_dict_index)
        print('相似案例个数:', len(cluster_menmbers_list[labels_last]) - 1)

        similar_case=[]
        for s in range(0, len(after_sort_dict_index)):
            print("第" + str(s) + "相似案例", fea_corpus[after_sort_dict_index[s]])
            similar_case.append(fea_corpus[after_sort_dict_index[s]])

        # TODO 连接Fuseki服务器。
        fuseki = JenaFuseki()
        # TODO 初始化自然语言到SPARQL查询的模块，参数是外部词典列表。
        q2s = Question2Sparql(
            ['./KB_query/external_dict/features.txt','./KB_query/external_dict/solutions.txt'])

        global fuzz1, fuzz_q1
        fuzz1, fuzz_q1 = [], []
        len_similar_case=len(similar_case)
        my_query=[]
        for i in range(1,len_similar_case):
            my_query.append(q2s.get_sparql(similar_case[i].encode('utf-8')))

        print('my_query',my_query)

        if my_query is not None:
            result_query=[]
            result_value=[]
            my_sol=[]
            valueq=[]
            for i in range(len(my_query)):
                result_query.append(fuseki.get_sparql_result(my_query[i]))
            for i in range(len(my_query)):
                result_value.append(fuseki.get_sparql_result_value(result_query[i]))


            print('result_query',result_query)
            print('result_value',result_value)

            for i in range(1,len(similar_case)):
                print("similar_case",similar_case)
                my_sol.append(similar_case[i])
                valueq.append(my_sol[i-1].split(' '))
            print('my_sol',my_sol)

            print('valueq',valueq)

            fuzz1=result_value
            fuzz_q1=valueq
            print('fuzz1',fuzz1)
            print('fuzzq1',fuzz_q1)
        else:
            # TODO 自然语言问题无法匹配到已有的正则模板上，回答“无法理解”
            print('I can\'t understand. :(')
        global  desc_text
        desc_text=[]
        desc_text.append("针对故障特征'"+str(w1)+str(w2)+str(w3)+str(w4)+str(w5)+str(w6)+str(w7)+str(w8)+f1+f2+f3+f4+f5+f6+f7+f8+f9+f10+"'有"+str(len(valueq))+"个相似案例")
        for i in range(1,len(valueq)+1):
            desc_text.append("案例"+str(i)+"：故障特征："+str(my_sol[i-1])+"解决方案特征："+str(fuzz1[i-1])+" 相似度："+str(after_sort_dict[i][1]))
        print("desc_text",desc_text)

    return render_template("fuzz_query.html")

#精确查询页面
@app.route("/precise_query",methods=['GET', 'POST'])
def precise():
    #初始化数据
    global f1, f2, f3, f4, f5, f6, list_fs, fuzz_q1, fuzz_q2, fuzz_q3
    global fuzz1, fuzz2, fuzz3
    global y1,y2,y3,y4,y5,y6,y7,y8,y9,y10,s1,s2,s3,s4,s5,s6,s7,s8,s9,s10,num1,num2,num3,num4,num5,num6,num7,num8,num9,num10,num11,num12,num13,num14,num15,num16,num17,num18
    y1=y2=y3=y4=y5=y6=y7=y8=y9=y10=None
    num1=num2=num3=num4=num5=num6=num7=num8=num9=num10=num11=num12=num13=num14=num15=num16=num17=num18=s1=""

    f1, f2, f3, f4, f5, f6 = None, None, None, None, None, None
    # 批量上传数据
    list_fs = []
    # 模糊查询
    fuzz_q1 = ['', '', '', '', '', '']
    fuzz_q2 = ['', '', '', '', '', '']
    fuzz_q3 = ['', '', '', '', '', '']
    fuzz1 = ['', '', '', '', '', '']
    fuzz2 = ['', '', '', '', '', '']
    fuzz3 = ['', '', '', '', '', '']

    return render_template("precise_query.html")


@app.route("/kg",methods=['GET', 'POST'])
def kg():
    return render_template("kg.html")

#精确查询
@app.route("/submit_precise",methods=['GET', 'POST'])
def submit_precise():
    print("submit_precise")
    global y1,y2,y3,y4,y5,y6,y7,y8,y9,y10,w1,w2,w3,w4,w5,w6,w7,w8,desc_text
    y1,y2,y3,y4,y5,y6,y7,y8,y9,y10,w1,w2,w3,w4,w5,w6,w7,w8=None,None,None,None,None,None,None,None,None,None,None,None,None,None,None,None,None,None
    desc_text=''
    y1 = request.form.get("fault_features1")
    y2 = request.form.get("fault_features2")
    y3 = request.form.get("fault_features3")
    y4 = request.form.get("fault_features4")
    y5 = request.form.get("fault_features5")
    y6 = request.form.get("fault_features6")
    y7 = request.form.get("fault_features7")
    y8 = request.form.get("fault_features8")
    y9 = request.form.get("fault_features9")
    y10 = request.form.get("fault_features10")
    w1=request.form.get("weather1")
    w2 = request.form.get("weather2")
    w3 = request.form.get("weather3")
    w4 = request.form.get("weather4")
    w5 = request.form.get("weather5")
    w6 = request.form.get("weather6")
    w7 = request.form.get("weather7")
    w8 = request.form.get("weather8")
    features=w1+w2+w3+w4+w5+w6+w7+w8+y1+y2+y3+y4+y5+y6+y7+y8+y9+y10
    y_list=[]
    if y1:
        y_list.append(y1)
    if y2:
        y_list.append(y2)
    if y3:
        y_list.append(y3)
    if y4:
        y_list.append(y4)
    if y5:
        y_list.append(y5)
    if y6:
        y_list.append(y6)
    if y7:
        y_list.append(y7)
    if y8:
        y_list.append(y8)
    if y9:
        y_list.append(y9)
    if y10:
        y_list.append(y10)

    print('y_list',y_list)


    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 passwd='zhangyixin',
                                 db='faults_last',
                                 port=3306,
                                 charset='utf8'
                                 )
    cur = connection.cursor()  # 游标（指针）cursor的方式操作数据
    sqlw1='SELECT COUNT(features) FROM features WHERE features="%s"' % w1
    sqlw2 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % w2
    sqlw3 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % w3
    sqlw4 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % w4
    sqlw5 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % w5
    sqlw6 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % w6
    sqlw7 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % w7
    sqlw8 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % w8
    sql1 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y1
    sql2 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y2
    sql3 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y3
    sql4 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y4
    sql5 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y5
    sql6 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y6
    sql7 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y7
    sql8 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y8
    sql9 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y9
    sql10 = 'SELECT COUNT(features) FROM features WHERE features="%s"' % y10

    cur.execute(sqlw1)
    see11 = cur.fetchone()
    cur.execute(sqlw2)
    see12 = cur.fetchone()
    cur.execute(sqlw3)
    see13 = cur.fetchone()
    cur.execute(sqlw4)
    see14 = cur.fetchone()
    cur.execute(sqlw5)
    see15 = cur.fetchone()
    cur.execute(sqlw6)
    see16 = cur.fetchone()
    cur.execute(sqlw7)
    see17 = cur.fetchone()
    cur.execute(sqlw8)
    see18 = cur.fetchone()


    cur.execute(sql1)  # execute(query, args):执行单条sql语句。
    see1 = cur.fetchone()  # 使结果全部可看
    cur.execute(sql2)
    see2 = cur.fetchone()
    cur.execute(sql3)
    see3 = cur.fetchone()
    cur.execute(sql4)
    see4= cur.fetchone()
    cur.execute(sql5)
    see5 = cur.fetchone()
    cur.execute(sql6)
    see6 = cur.fetchone()
    cur.execute(sql7)
    see7 = cur.fetchone()
    cur.execute(sql8)
    see8 = cur.fetchone()
    cur.execute(sql9)
    see9 = cur.fetchone()
    cur.execute(sql10)
    see10 = cur.fetchone()

    global num1,num2,num3,num4,num5,num6,num7,num8,num9,num10,num11,num12,num13,num14,num15,num16,num17,num18
    data = []
    data.append(see11)
    data.append(see12)
    data.append(see13)
    data.append(see14)
    data.append(see15)
    data.append(see16)
    data.append(see17)
    data.append(see18)
    data.append(see1)
    data.append(see2)
    data.append(see3)
    data.append(see4)
    data.append(see5)
    data.append(see6)
    data.append(see7)
    data.append(see8)
    data.append(see9)
    data.append(see10)
    print('data',data)
    num1=data[0][0]
    num2=data[1][0]
    num3 = data[2][0]
    num4 = data[3][0]
    num5 = data[4][0]
    num6 = data[5][0]
    num7 = data[6][0]
    num8 = data[7][0]
    num9 = data[8][0]
    num10 = data[9][0]
    num11=data[10][0]
    num12 = data[11][0]
    num13 = data[12][0]
    num14 = data[13][0]
    num15 = data[14][0]
    num16 = data[15][0]
    num17 = data[16][0]
    num18 = data[17][0]

    connection.commit()
    connection.close()


    # TODO 连接Fuseki服务器。
    fuseki = JenaFuseki()

    global solution,s1

    SPARQL_PREXIX = u"""
                                        PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
                                        PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
                                        PREFIX : <http://www.kgdemo.com#>
                                        """
    SPARQL_SELECT_TEM = u"{prefix}\n" + \
                        u"SELECT DISTINCT {select} WHERE {{\n" + \
                        u"{expression}\n" + \
                        u"}}\n"
    select = u"?x"
    fuseki = JenaFuseki()
    if len(y_list)==1:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'."\
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1,input2=w2,input3=w3,input4=w4,input5=w5,input6=w6,input7=w7,input8=w8,input9=y1)
    elif len(y_list)==2:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p10 :features '{input10}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?p10 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5, input6=w6, input7=w7,
                                          input8=w8, input9=y1,input10=y2)
    elif len(y_list)==3:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p10 :features '{input10}'." \
            u"?p11 :features '{input11}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?p10 :fea2sol ?m." \
            u"?p11 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5, input6=w6, input7=w7,
                                          input8=w8, input9=y1,input10=y2,input11=y3)
    elif len(y_list)==4:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p10 :features '{input10}'." \
            u"?p11 :features '{input11}'." \
            u"?p12 :features '{input12}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?p10 :fea2sol ?m." \
            u"?p11 :fea2sol ?m." \
            u"?p12 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5, input6=w6, input7=w7,
                                          input8=w8, input9=y1,input10=y2,input11=y3,input12=y4),
    elif len(y_list)==5:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p10 :features '{input10}'." \
            u"?p11 :features '{input11}'." \
            u"?p12 :features '{input12}'." \
            u"?p13 :features '{input13}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?p10 :fea2sol ?m." \
            u"?p11 :fea2sol ?m." \
            u"?p12 :fea2sol ?m." \
            u"?p13 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5, input6=w6, input7=w7,
                                          input8=w8, input9=y1,input10=y2,input11=y3,input12=y4,input13=y5)
    elif len(y_list)==6:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p10 :features '{input10}'." \
            u"?p11 :features '{input11}'." \
            u"?p12 :features '{input12}'." \
            u"?p13 :features '{input13}'." \
            u"?p14 :features '{input14}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?p10 :fea2sol ?m." \
            u"?p11 :fea2sol ?m." \
            u"?p12 :fea2sol ?m." \
            u"?p13 :fea2sol ?m." \
            u"?p14 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5, input6=w6, input7=w7,
                                          input8=w8, input9=y1,input10=y2,input11=y3,input12=y4,input13=y5,input14=y6)
    elif len(y_list)==7:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p10 :features '{input10}'." \
            u"?p11 :features '{input11}'." \
            u"?p12 :features '{input12}'." \
            u"?p13 :features '{input13}'." \
            u"?p14 :features '{input14}'." \
            u"?p15 :features '{input15}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?p10 :fea2sol ?m." \
            u"?p11 :fea2sol ?m." \
            u"?p12 :fea2sol ?m." \
            u"?p13 :fea2sol ?m." \
            u"?p14 :fea2sol ?m." \
            u"?p15 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5, input6=w6, input7=w7,
                                          input8=w8, input9=y1,input10=y2,input11=y3,input12=y4,input13=y5,input14=y6,input15=y7)
    elif len(y_list)==8:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p10 :features '{input10}'." \
            u"?p11 :features '{input11}'." \
            u"?p12 :features '{input12}'." \
            u"?p13 :features '{input13}'." \
            u"?p14 :features '{input14}'." \
            u"?p15 :features '{input15}'." \
            u"?p16 :features '{input16}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?p10 :fea2sol ?m." \
            u"?p11 :fea2sol ?m." \
            u"?p12 :fea2sol ?m." \
            u"?p13 :fea2sol ?m." \
            u"?p14 :fea2sol ?m." \
            u"?p15 :fea2sol ?m." \
            u"?p16 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5, input6=w6, input7=w7,
                                          input8=w8, input9=y1,input10=y2,input11=y3,input12=y4,input13=y5,input14=y6,input15=y7,input16=y8)
    elif len(y_list)==9:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p10 :features '{input10}'." \
            u"?p11 :features '{input11}'." \
            u"?p12 :features '{input12}'." \
            u"?p13 :features '{input13}'." \
            u"?p14 :features '{input14}'." \
            u"?p15 :features '{input15}'." \
            u"?p16 :features '{input16}'." \
            u"?p17 :features '{input17}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?p10 :fea2sol ?m." \
            u"?p11 :fea2sol ?m." \
            u"?p12 :fea2sol ?m." \
            u"?p13 :fea2sol ?m." \
            u"?p14 :fea2sol ?m." \
            u"?p15 :fea2sol ?m." \
            u"?p16 :fea2sol ?m." \
            u"?p17 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5, input6=w6, input7=w7,
                                          input8=w8, input9=y1,input10=y2,input11=y3,input12=y4,input13=y5,input14=y6,input15=y7,input16=y8,input17=y9)
    elif len(y_list)==10:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p6 :features '{input6}'." \
            u"?p7 :features '{input7}'." \
            u"?p8 :features '{input8}'." \
            u"?p9 :features '{input9}'." \
            u"?p10 :features '{input10}'." \
            u"?p11 :features '{input11}'." \
            u"?p12 :features '{input12}'." \
            u"?p13 :features '{input13}'." \
            u"?p14 :features '{input14}'." \
            u"?p15 :features '{input15}'." \
            u"?p16 :features '{input16}'." \
            u"?p17 :features '{input17}'." \
            u"?p18 :features '{input18}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?p6 :fea2sol ?m." \
            u"?p7 :fea2sol ?m." \
            u"?p8 :fea2sol ?m." \
            u"?p9 :fea2sol ?m." \
            u"?p10 :fea2sol ?m." \
            u"?p11 :fea2sol ?m." \
            u"?p12 :fea2sol ?m." \
            u"?p13 :fea2sol ?m." \
            u"?p14 :fea2sol ?m." \
            u"?p15 :fea2sol ?m." \
            u"?p16 :fea2sol ?m." \
            u"?p17 :fea2sol ?m." \
            u"?p18 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5, input6=w6, input7=w7,
                                          input8=w8, input9=y1,input10=y2,input11=y3,input12=y4,input13=y5,input14=y6,
                                          input15=y7,input16=y8,input17=y9,input18=y10)
    else:
        e = u"?p1 :features '{input1}'." \
            u"?p2 :features '{input2}'." \
            u"?p3 :features '{input3}'." \
            u"?p4 :features '{input4}'." \
            u"?p5 :features '{input5}'." \
            u"?p1 :fea2sol ?m." \
            u"?p2 :fea2sol ?m." \
            u"?p3 :fea2sol ?m." \
            u"?p4 :fea2sol ?m." \
            u"?p5 :fea2sol ?m." \
            u"?m :sol_features ?x".format(input1=w1, input2=w2, input3=w3, input4=w4, input5=w5)

        print("出错")
        s1=''
    my_query = SPARQL_SELECT_TEM.format(prefix=SPARQL_PREXIX,
                                        select=select,
                                        expression=e)
    print('my_query', my_query)

    result = fuseki.get_sparql_result(my_query)
    print('result:', result)
    value = fuseki.get_sparql_result_value(result)
    print('value:', value)


    if value!=[]:
        s1=value[0]
    else:
        print('没有查到此案例')
        s1=''

        # TODO 判断结果是否是布尔值，是布尔值则提问类型是"ASK"，回答“是”或者“不知道”。
        # if isinstance(value, bool):
        #     if value is True:
        #         print('Yes')
        #     else:
        #         print('I don\'t know. :(')
        # else:
        #     # TODO 查询结果为空，根据OWA，回答“不知道”
        #     s1=value[0]


    global prec_text
    prec_text = []
    print(y1 + "&")
    if s1 == '':
        prec_text.append("未查询到相关案例！")
    else:
        prec_text.append(
            "针对故障特征：" +w1+" "+w2+" "+w3+" "+w4+" "+w5+" "+w6+" "+w7+" "+w8+" "+ y1 + " " + y2 + " " + y3 + " " + y4 + " " + y5 + " " + y6 + " " + y7 + " " + y8 + " " + y9 + " " + y10)
        prec_text.append(
            "解决方案特征为：" + s1 + " " )
    print('prec_text',prec_text)

    return render_template("precise_query.html",t1=y1,t2=y2,t3=y3,t4=y4,t5=y5,t6=y6)


#饼状图数据
@app.route("/barData")
def get_bar_chart():
    global num1,num2,num3,num4,num5,num6,num7,num8,num9,num10,y1,y2,y3,y4,y5,y6,y7,y8,y9,y10
    global w1,w2,w3,w4,w5,w6,w7,w8
    global num11,num12,num13,num14,num15,num16,num17,num18
    value = [num1,num2,num3,num4,num5,num6,num7,num8,num9,num10,num11,num12,num13,num14,num15,num16,num17,num18]
    name=[w1,w2,w3,w4,w5,w6,w7,w8,y1,y2,y3,y4,y5,y6,y7,y8,y9,y10]
    try:
        if y10:
            name=[w1,w2,w3,w4,w5,w6,w7,w8,y1,y2,y3,y4,y5,y6,y7,y8,y9,y10]
        elif y9:
            name = [w1,w2,w3,w4,w5,w6,w7,w8,y1, y2, y3, y4, y5, y6, y7, y8, y9]
        elif y8:
            name = [w1,w2,w3,w4,w5,w6,w7,w8,y1, y2, y3, y4, y5, y6, y7, y8]
        elif y7:
            name = [w1,w2,w3,w4,w5,w6,w7,w8,y1, y2, y3, y4, y5, y6, y7]
        elif y6:
            name = [w1,w2,w3,w4,w5,w6,w7,w8,y1, y2, y3, y4, y5, y6]
        elif y5:
            name = [w1,w2,w3,w4,w5,w6,w7,w8,y1, y2, y3, y4, y5]
        elif y4:
            name = [w1,w2,w3,w4,w5,w6,w7,w8,y1, y2, y3, y4]
        elif y3:
            name = [w1,w2,w3,w4,w5,w6,w7,w8,y1, y2, y3]
        elif y2:
            name = [w1,w2,w3,w4,w5,w6,w7,w8,y1, y2]
    except BaseException as e:
        print(e)
    else:
        print(value)
    # if num6!='':
    #     name=["feature1","feature2","feature3","feature4","feature5","feature6"]
    # elif num5!='':
    #     name = ["feature1", "feature2", "feature3", "feature4", "feature5"]
    # elif num4!='':
    #     name = ["feature1", "feature2", "feature3", "feature4"]
    # else:
    #     name = ["feature1", "feature2", "feature3"]
    return json.dumps({'name':name,'value':value})

#精确查询数据
@app.route("/linkData",methods=['GET', 'POST'])
def get_linkData():
    print("*************")
    global y1,y2,y3,y4,y5,y6,y7,y8,y9,y10,w1,w2,w3,w4,w5,w6,w7,w8
    global s1
    global prec_text


    list=[]
    for i in range(1):
        list.append({"name": '案例',"des": '', "symbolSize": 50, "category": 0,}, )
        list.append({"name": '特征' ,"des": '', "symbolSize": 50, "category":0, }, )
        list.append({"name": '解决方案' ,"des": s1, "symbolSize": 50, "category": 0, }, )
        # list.append({"name": y1, "des": 'nodedes01', "symbolSize": 50, "category": 0, }, )
        # list.append({"name": y2, "des": 'nodedes01', "symbolSize": 50, "category": 0, }, )
        # list.append({"name": s1, "des": 'nodedes01', "symbolSize": 50, "category": 0, }, )
        # list.append({"name": s2, "des": 'nodedes01', "symbolSize": 50, "category": 0, }, )
    links=[
    ]
    for i in range(1):
        links.append({"source": '案例', "target": "特征", "name": ' '},)
        links.append({"source": '案例', "target": "解决方案", "name": ' '},)
        links.append({"source": '特征' , "target": y1 ,"name":' '}, )
        links.append({"source": '特征', "target": y2 ,"name":' '}, )
        # links.append({"source": '解决方案', "target": s1,"name":'特征1'}, )
        # links.append({"source": '解决方案', "target": s2, "name": '特征2'}, )
    if y3:
        list.append({"name": y3, "des": '特征3', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":y3,"name":' '},)

    if y4:
        list.append({"name": y4, "des": '特征4', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":y4,"name":' '},)

    if y5:
        list.append({"name": y5, "des": '特征5', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":y5,"name":' '},)

    if y6:
        list.append({"name": y6, "des": '特征6', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":y6,"name":' '},)

    if y7:
        list.append({"name": y7, "des": '特征7', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":y7,"name":' '},)

    if y8:
        list.append({"name": y8, "des": '特征8', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":y8,"name":' '},)

    if y9:
        list.append({"name": y9, "des": '特征9', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":y9,"name":' '},)

    if y10:
        list.append({"name": y10, "des": '特征10', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":y10,"name":' '},)

    if w1:
        list.append({"name": w1, "des": '车辆类型', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":w1,"name":' '},)
    if w2:
        list.append({"name": w2, "des": '车辆型号', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":w2,"name":' '},)
    if w3:
        list.append({"name": w3, "des": '温度', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":w3,"name":' '},)
    if w4:
        list.append({"name": w4, "des": '湿度', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":w4,"name":' '},)
    if w5:
        list.append({"name": w5, "des": '气压', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":w5,"name":' '},)
    if w6:
        list.append({"name": w6, "des": '海拔', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":w6,"name":' '},)
    if w7:
        list.append({"name": w7, "des": '盐度', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":w7,"name":' '},)
    if w8:
        list.append({"name": w8, "des": '昼夜温差', "symbolSize": 50, "category": 0, }, )
        links.append({"source": '特征', "target":w8,"name":' '},)

    return json.dumps({'list':list,'links':links,'prec_text':prec_text})

#模糊查询数据
@app.route("/fuzzData",methods=['GET', 'POST'])
def get_fuzzData():
    print("fuzzData")
    global fuzz1,fuzz_q1,v1
    global f1,f2,f3,f4,f5,f6,f7,f8,f9,f10,w1,w2,w3,w4,w5,w6,w7,w8
    fuzz_features=""
    print('f1',f1)
    print('f2',f2)
    print('w1',w1)
    print('w2',w2)
    if w1 !='':
        fuzz_features+=w1+"&"
    if w2 != '':
        fuzz_features += w2 + "&"
    if w3 !='':
        fuzz_features+=w3+"&"
    if w4 != '':
        fuzz_features += w4 + "&"
    if w5 !='':
        fuzz_features+=w5+"&"
    if w6 != '':
        fuzz_features += w6 + "&"
    if w7 !='':
        fuzz_features+=w7+"&"
    if w8 != '':
        fuzz_features += w8 + "&"

    if f1 !='':
        fuzz_features+=f1+"&"
    if f2 !='':
        fuzz_features += f2 + "&"
    if f3 !='':
        fuzz_features += f3 + "&"
    if f4 !='':
        fuzz_features += f4 + "&"
    if f5 !='':
        fuzz_features += f5 + "&"
    if f6 !='':
        fuzz_features += f6 + "&"
    if f7 !='':
        fuzz_features += f7 + "&"
    if f8 !='':
        fuzz_features += f8 + "&"
    if f9 !='':
        fuzz_features += f9 + "&"
    if f10 !='':
        fuzz_features += f10
    print("AAAAAAAAAAAAAAAAAAAAAAAAa")

    list=[
        {"name": '特征', "des": '特征' + fuzz_features, "symbolSize": 50, "category": 0, },
    ]
    links=[
    ]
    fuzz_len=len(fuzz1)
    print('fuzz_len',fuzz_len)
    number=int(v1)
    if number<fuzz_len:
        for i in range(number):
            list.append({"name":'案例'+str(i+1), "des": '案例'+str(i+1), "symbolSize": 50, "category": 0, }, )
            list.append({"name": str(i+1)+'特征', "des": str(i+1)+'特征', "symbolSize": 50, "category": 1, }, )
            # list.append({"name": str(i+1)+'解决方案', "des": '', "symbolSize": 50, "category": 2, }, )
            links.append({"source": '特征', "target": '案例'+str(i+1),"name":' ' },)
            links.append({"source": '案例'+str(i+1), "target": str(i+1)+'特征',"name":' '},)
            links.append({"source": '案例'+str(i+1), "target": str(i+1)+'解决方案',"name":' '},)
        for j in range(0,number):
            print("YYYYYYYYYYYY")
            print(fuzz1[j][0])
            list.append({"name": str(j + 1) + '解决方案', "des": fuzz1[j][0], "symbolSize": 50, "category": 2, }, )
                # list.append({"name": fuzz1[j][k], "des": 'nodedes01', "symbolSize": 50, "category": 2, }, )
                # links.append({"source": str(j+1)+'解决方案', "target": fuzz1[j][k], "name": '特征' + str(k + 1)},)
        for a in range(number):
            for b in range(len(fuzz_q1[a])):
                list.append({"name": fuzz_q1[a][b], "des": fuzz_q1[a][b], "symbolSize": 50, "category": 1, }, )
                links.append({"source":str(a+1)+ '特征', "target": fuzz_q1[a][b],"name":'特征'+str(b+1)},)
    else:
        for i in range(fuzz_len):
            list.append({"name": '案例' + str(i + 1), "des":  '案例' + str(i + 1), "symbolSize": 50, "category": 0, }, )
            list.append({"name": str(i + 1) + '特征', "des": str(i + 1) + '特征', "symbolSize": 50, "category": 1, }, )
            # list.append({"name": str(i + 1) + '解决方案', "des": 'nodedes01', "symbolSize": 50, "category": 2, }, )
            links.append({"source": '特征', "target": '案例' + str(i + 1), "name": ' '}, )
            links.append({"source": '案例' + str(i + 1), "target": str(i + 1) + '特征', "name": ' '}, )
            links.append({"source": '案例' + str(i + 1), "target": str(i + 1) + '解决方案', "name": ' '}, )
        for j in range(fuzz_len):
            list.append({"name": str(j + 1) + '解决方案', "des": fuzz1[j][0], "symbolSize": 50, "category": 2, }, )
                # list.append({"name": fuzz1[j][k], "des": 'nodedes01', "symbolSize": 50, "category": 2, }, )
                # links.append({"source": str(j + 1) + '解决方案', "target": fuzz1[j][k], 'name': '特征' + str(k + 1)}, )
        for a in range(fuzz_len):
            for b in range(len(fuzz_q1[a])):
                list.append({"name": fuzz_q1[a][b], "des": fuzz_q1[a][b], "symbolSize": 50, "category": 1, }, )
                links.append({"source": str(a + 1) + '特征', "target": fuzz_q1[a][b],'name':'特征'+str(b+1)}, )
    print("list:",list)
    print("links:",links)
    global desc_text
    return json.dumps({'list':list,'links': links,'desc_text':desc_text})

#知识地图页面
@app.route("/map", methods=['GET', 'POST'])
def map1():
    return render_template("map.html")

#整体知识地图数据
@app.route("/allData",methods=['GET', 'POST'])
def get_allData():
    if request.method == 'POST':
        global list_map, links_map
        list = []
        links = []
        list_map = []
        links_map = []
        # if request.method=='POST':
        v1 = request.form.get("v1")
        t1 = request.form.get("t1")
        input1 = request.form.get("input1")
        print("v1:",v1)
        print("t1:", t1)
        print("input1:",input1)

        if (v1 == '1'):
            name=input1
            SPARQL_PREXIX = u"""
                        PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
                        PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
                        PREFIX : <http://www.kgdemo.com#>
                        """
            SPARQL_SELECT_TEM = u"{prefix}\n" + \
                                u"SELECT DISTINCT {select} WHERE {{\n" + \
                                u"{expression}\n" + \
                                u"}}\n"
            select = u"?x"
            fuseki = JenaFuseki()
            e = u"?s :group_id '{input1}'." \
                u"?s :features ?x.".format(input1=input1)
            my_query = SPARQL_SELECT_TEM.format(prefix=SPARQL_PREXIX,
                                                select=select,
                                                expression=e)
            print('my_query',my_query)

            e1 = u"?s :sol_group_id '{input1}'." \
                u"?s :sol_features ?x.".format(input1=input1)
            my_query1 = SPARQL_SELECT_TEM.format(prefix=SPARQL_PREXIX,
                                                select=select,
                                                expression=e1)

            result = fuseki.get_sparql_result(my_query)
            result1 = fuseki.get_sparql_result(my_query1)
            print('result:', result)
            print('result1:', result1)
            value = fuseki.get_sparql_result_value(result)
            value1 = fuseki.get_sparql_result_value(result1)
            print('value:', value)
            print('value1:', value1)
            length = len(value)
            length1 = len(value1)


            list.append({"name": '案例'+str(name), "des": '', "symbolSize": 50, "category": 0, }, )
            list.append({"name": '故障特征', "des": '', "symbolSize": 50, "category": 0, }, )
            # list.append({"name": '解决方案特征' , "des": 'nodedes01', "symbolSize": 50, "category": 0, }, )
            links.append({"source":  '案例'+str(name), "target": '故障特征', "name": ''}, )
            links.append({"source":  '案例'+str(name), "target": '解决方案特征', "name": ''}, )

            color_num=2
            list_map.append([])
            for i in range(0, length):
                list.append({"name": value[i], "des": value[i], "symbolSize": 40, "category": 1, }, )
                links.append({"source": '故障特征', "target": value[i] , "name": '特征'+str(i+1)}, )
                list_map[0].append(value[i])

            list_map.append([])
            for j in range(0, length1):
                list.append({"name": '解决方案特征', "des":  value1[j], "symbolSize": 40, "category": 1, }, )
                # links.append({"source": '解决方案特征', "target": value1[j] , "name": '特征 '+str(j+1)}, )
                list_map[1].append(value1[j])

            print(links)
            print(list)
            print('color_num',color_num)
            return json.dumps({'list': list, 'links': links,'color_num':color_num})
        elif(v1=='2'):
            name = input1
            SPARQL_PREXIX = u"""
                                    PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
                                    PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
                                    PREFIX : <http://www.kgdemo.com#>
                                    """
            SPARQL_SELECT_TEM = u"{prefix}\n" + \
                                u"SELECT DISTINCT {select} WHERE {{\n" + \
                                u"{expression}\n" + \
                                u"}}\n"
            select = u"?x"
            fuseki = JenaFuseki()
            e = u"?s :features '{input1}'." \
                u"?s :group_id ?x.".format(input1=input1)
            my_query = SPARQL_SELECT_TEM.format(prefix=SPARQL_PREXIX,
                                                select=select,
                                                expression=e)
            print('my_query', my_query)

            result = fuseki.get_sparql_result(my_query)
            print('result:', result)
            value = fuseki.get_sparql_result_value(result)
            print('value:', value)
            length = len(value)

            list.append({"name": name, "des": name, "symbolSize": 50, "category": 0, }, )
            color_num = length + 1
            list_map.append([])
            for i in range(0, length):
                list.append({"name": '案例' + value[i], "des": '案例' + value[i], "symbolSize": 40, "category": i + 1, }, )
                links.append({"source": name, "target": '案例' + value[i], "name": '案例 ' + str(i+1)}, )
                list_map[0].append(value[i])
            print('list_map',list_map)
            print('links',links)
            print('list',list)
            print('color_num', color_num)
            return json.dumps({'list': list, 'links': links, 'color_num': color_num})

        elif(v1=='3'):
            name = input1
            SPARQL_PREXIX = u"""
                                    PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
                                    PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
                                    PREFIX : <http://www.kgdemo.com#>
                                    """
            SPARQL_SELECT_TEM = u"{prefix}\n" + \
                                u"SELECT DISTINCT {select} WHERE {{\n" + \
                                u"{expression}\n" + \
                                u"}}\n"
            select = u"?x"
            fuseki = JenaFuseki()
            e = u"?s :sol_features '{input1}'." \
                u"?s :sol_group_id ?x.".format(input1=input1)
            my_query = SPARQL_SELECT_TEM.format(prefix=SPARQL_PREXIX,
                                                select=select,
                                                expression=e)
            print('my_query', my_query)

            result = fuseki.get_sparql_result(my_query)
            print('result:', result)
            value = fuseki.get_sparql_result_value(result)
            print('value:', value)
            length = len(value)

            list.append({"name": name, "des": name, "symbolSize": 50, "category": 0, }, )
            color_num = length + 1
            list_map.append([])
            for i in range(0, length):
                list.append({"name": '案例' + value[i], "des": '案例' + value[i], "symbolSize": 40, "category": i + 1, }, )
                links.append({"source": name, "target": '案例' + value[i], "name": '案例 ' + str(i)}, )
                list_map[0].append(value[i])
            print('list_map',list_map)
            print('links',links)
            print('list',list)
            print('color_num', color_num)
            return json.dumps({'list': list, 'links': links, 'color_num': color_num})
        else:
            SPARQL_PREXIX = u"""
                PREFIX rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#>
                PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
                PREFIX : <http://www.kgdemo.com#>
                """

            SPARQL_SELECT_TEM = u"{prefix}\n" + \
                                u"SELECT DISTINCT {select} WHERE {{\n" + \
                                u"{expression}\n" + \
                                u"}}\n"
            select = u"?x"
            connection = pymysql.connect(host='localhost',
                                         user='root',
                                         passwd='zhangyixin',
                                         db='faults_last',
                                         port=3306,
                                         charset='utf8'
                                         )
            cur = connection.cursor()  # 游标（指针）cursor的方式操作数据
            sql1 = 'SELECT count(*) FROM(SELECT count(*) FROM features GROUP BY features.group_id)a'
            cur.execute(sql1)  # execute(query, args):执行单条sql语句。
            see1 = cur.fetchone()  # 使结果全部可看
            print(see1[0])
            connection.commit()
            connection.close()
            fuseki = JenaFuseki()
            list = []
            links = []
            children=[]

            color_num=see1[0] + 1

            # for i in range(1,see1[0]+1):
            #     list.append({"name": '案例'+str(i), "value": 0, "children": [{"name":'故障特征',"value":1,"children":[]},{"name":'解决方案',"value":1,"children":[]}]}, )
            for i in range(1, see1[0] + 1):
                list.append({"name": '案例' + str(i), "des":  '', "symbolSize": 50, "category": i, }, )
                list.append({"name": str(i) + '特征', "des": '', "symbolSize": 50, "category": i, }, )
                # list.append({"name": str(i) + '故障方案', "des": 'nodedes01', "symbolSize": 50, "category": i, }, )
                links.append({"source": '案例' + str(i), "target": str(i) + '特征', "name": ' '}, )
                links.append({"source": '案例' + str(i), "target": str(i) + '故障方案', "name": ' '}, )
            for i in range(1, see1[0] + 1):
                e = u"?s :features ?x." \
                    u"?s :group_id '{a}'.".format(a=i)
                my_query = SPARQL_SELECT_TEM.format(prefix=SPARQL_PREXIX,
                                                    select=select,
                                                    expression=e)
                result = fuseki.get_sparql_result(my_query)
                value = fuseki.get_sparql_result_value(result)
                length = len(value)

                # for j in range(1, length + 1):
                #     children.append({"name": value[j - 1], "value": 0, "children": []}, )

                for j in range(0, length ):
                    list_map.append([])
                    list.append({"name": value[j ], "des": value[j ], "symbolSize": 50, "category": j+1, }, )
                    links.append({"source": str(i+1) + '特征', "target": value[j ], 'name': '特征' + str(j+1)}, )
                    list_map[j].append(value[j])
            for i in range(1, see1[0] + 1):
                e = u"?s :sol_group_id '{b}'." \
                    u" ?s :sol_features ?x.".format(b=i)
                my_query1 = SPARQL_SELECT_TEM.format(prefix=SPARQL_PREXIX,
                                                     select=select,
                                                     expression=e)
                result1 = fuseki.get_sparql_result(my_query1)
                value1 = fuseki.get_sparql_result_value(result1)
                length1 = len(value1)
                print('length1',length1)
                print('value1',value1)

                for j in range(0, length1 ):
                    links_map.append([])
                    list.append({"name": str(i+1) + '故障方案', "des": value1[j], "symbolSize": 50, "category": j, }, )
                    links_map[j].append(value1[j])

                # for j in range(1, length1 + 1):
                #     list.append({"name": value1[j - 1], "des": 'nodedes01', "symbolSize": 50, "category": j, }, )
                #     links.append({"source": str(i) + '故障方案', "target": value1[j - 1], 'name': '特征' + str(j)}, )
    print('list_map',list_map)
    print('links_map',links_map)
    print('links',links)
    print('list',list)
    print('color_num',color_num)
    return json.dumps({'list': list, 'links': links,'color_num':color_num})

#整体知识图谱导出
@app.route("/output",methods=['GET', 'POST'])
def output():
    if request.method == 'POST':
        global links_map,list_map
        print("list_map",list_map)
        print("links_map",links_map)
        workbook = xlsxwriter.Workbook('故障.xlsx')  # 创建一个名为 hello.xlsx 赋值给workbook
        worksheet = workbook.add_worksheet()  # 创建一个默认工作簿 赋值给worksheet
        # 工作簿也支持命名，
        # 如：workbook.add_worksheet('hello')
        alpha=['A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z']
        if len(list_map)==2:
            # worksheet.write('A1', 'Hello ld')
            worksheet.write('A1', '故障特征')
            worksheet.write('A2', '车辆类型')
            worksheet.write('A3', '车辆型号')
            worksheet.write('A4', '温度')
            worksheet.write('A5', '湿度')
            worksheet.write('A6', '气压')
            worksheet.write('A7', '海拔')
            worksheet.write('A8', '盐度')
            worksheet.write('A9', '昼夜温差')
            worksheet.write('B1', '故障特征值')
            worksheet.write('B2', list_map[0][6])
            worksheet.write('B3', list_map[0][7])
            worksheet.write('B4', list_map[0][0])
            worksheet.write('B5', list_map[0][1])
            worksheet.write('B6', list_map[0][2])
            worksheet.write('B7', list_map[0][3])
            worksheet.write('B8', list_map[0][4])
            worksheet.write('B9', list_map[0][5])
            for i in range(10,10+len(list_map[0])-8):
                worksheet.write('A'+str(i), '故障特征'+str(i-9))
            for i in range(10,10+len(list_map[0])-8):
                worksheet.write('B'+str(i),list_map[0][i-2] )

            worksheet.write('A'+str(i+1), '解决方案')
            worksheet.write('B' + str(i+1) , list_map[1][0])
            # for i in range(0,len(list_map[0])):
            #     worksheet.write(alpha[i+1]+'2', list_map[0][i])
            # for i in range(0,len(list_map[1])):
            #     worksheet.write(alpha[i+11]+'2', list_map[1][i])
        elif len(list_map)==1:
            print("故障特征查询或解决方案查询")
            print(len(list_map))
            print(len(list_map[0]))
            print(len(list_map[0][0]))
            if len(list_map[0][0])>1:
                print("解决方案查询")
            else:
                print("故障特征查询")

        else:
            print("整体知识地图")
        workbook.close()  # 关闭工作簿
        return render_template("map.html")


#精确查询导出
@app.route("/output_precise",methods=['GET', 'POST'])
def output_precise():
    if request.method == 'POST':
        global y1, y2, y3, y4, y5, y6, y7, y8, y9, y10, w1, w2, w3, w4, w5, w6, w7, w8
        global s1
        y_list=[]
        if y1:
            y_list.append(y1)
        if y2:
            y_list.append(y2)
        if y3:
            y_list.append(y3)
        if y4:
            y_list.append(y4)
        if y5:
            y_list.append(y5)
        if y6:
            y_list.append(y6)
        if y7:
            y_list.append(y7)
        if y8:
            y_list.append(y8)
        if y9:
            y_list.append(y9)
        if y10:
            y_list.append(y10)
        print('y_list',y_list)


        workbook = xlsxwriter.Workbook('精确查询故障导出.xlsx')  # 创建一个名为 hello.xlsx 赋值给workbook
        worksheet = workbook.add_worksheet()  # 创建一个默认工作簿 赋值给worksheet
        # 工作簿也支持命名，
        # 如：workbook.add_worksheet('hello')
        worksheet.write('A1', '故障特征')
        worksheet.write('A2', '车辆类型')
        worksheet.write('A3', '车辆型号')
        worksheet.write('A4', '温度')
        worksheet.write('A5', '湿度')
        worksheet.write('A6', '气压')
        worksheet.write('A7', '海拔')
        worksheet.write('A8', '盐度')
        worksheet.write('A9', '昼夜温差')
        worksheet.write('B1', '故障特征值')
        worksheet.write('B2', w1)
        worksheet.write('B3', w2)
        worksheet.write('B4', w3)
        worksheet.write('B5', w4)
        worksheet.write('B6', w5)
        worksheet.write('B7', w6)
        worksheet.write('B8', w7)
        worksheet.write('B9', w8)
        for i in range(10,10+len(y_list)):
            worksheet.write('A'+str(i), '故障特征'+str(i-9))
        for i in range(10,10+len(y_list)):
            worksheet.write('B'+str(i),y_list[i-10] )

        worksheet.write('A'+str(i+1), '解决方案')
        worksheet.write('B' + str(i+1) , s1)


        workbook.close()  # 关闭工作簿
        return render_template("precise_query.html")

#模糊查询导出
@app.route("/output_fuzz",methods=['GET', 'POST'])
def output_fuzz():
    if request.method == 'POST':
        global fuzz1, fuzz_q1, v1
        global f1, f2, f3, f4, f5, f6, f7, f8, f9, f10, w1, w2, w3, w4, w5, w6, w7, w8

        workbook = xlsxwriter.Workbook('模糊查询故障导出.xlsx')  # 创建一个名为 hello.xlsx 赋值给workbook
        worksheet = workbook.add_worksheet()  # 创建一个默认工作簿 赋值给worksheet
        # 工作簿也支持命名，
        # 如：workbook.add_worksheet('hello')
        alpha=['A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z']

        # worksheet.write('A1', 'Hello ld')
        worksheet.write('A1', '序号')
        worksheet.write('B1', '车辆类型')
        worksheet.write('C1', '车辆型号')
        worksheet.write('D1', '温度')
        worksheet.write('E1', '湿度')
        worksheet.write('F1', '气压')
        worksheet.write('G1', '海拔')
        worksheet.write('H1', '盐度')
        worksheet.write('I1', '昼夜温差')
        worksheet.write('J1', '故障特征1')
        worksheet.write('K1', '故障特征2')
        worksheet.write('L1', '故障特征3')
        worksheet.write('M1', '故障特征4')
        worksheet.write('N1', '故障特征5')
        worksheet.write('O1', '故障特征6')
        worksheet.write('P1', '故障特征7')
        worksheet.write('Q1', '故障特征8')
        worksheet.write('R1', '故障特征9')
        worksheet.write('S1', '故障特征10')
        worksheet.write('T1', '解决方案')

        print('fuzz1',fuzz1)
        print('fuzz_q1',fuzz_q1)
        for i in range(0,int(v1)):
            worksheet.write(i + 1, 0, i + 1)
            worksheet.write(i + 1, 1, fuzz_q1[i][6])
            worksheet.write(i + 1, 2, fuzz_q1[i][7])
            worksheet.write(i + 1, 3, fuzz_q1[i][0])
            worksheet.write(i + 1, 4, fuzz_q1[i][1])
            worksheet.write(i + 1, 5, fuzz_q1[i][2])
            worksheet.write(i + 1, 6, fuzz_q1[i][3])
            worksheet.write(i + 1, 7, fuzz_q1[i][4])
            worksheet.write(i + 1, 8, fuzz_q1[i][5])
            worksheet.write(i + 1, 19, fuzz1[i][0])
            if len(fuzz_q1[i])==9:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
            elif len(fuzz_q1[i])==10:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
                worksheet.write(i + 1, 10, fuzz_q1[i][9])
            elif len(fuzz_q1[i])==11:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
                worksheet.write(i + 1, 10, fuzz_q1[i][9])
                worksheet.write(i + 1, 11, fuzz_q1[i][10])
            elif len(fuzz_q1[i])==12:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
                worksheet.write(i + 1, 10, fuzz_q1[i][9])
                worksheet.write(i + 1, 11, fuzz_q1[i][10])
                worksheet.write(i + 1, 12, fuzz_q1[i][11])
            elif len(fuzz_q1[i])==13:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
                worksheet.write(i + 1, 10, fuzz_q1[i][9])
                worksheet.write(i + 1, 11, fuzz_q1[i][10])
                worksheet.write(i + 1, 12, fuzz_q1[i][11])
                worksheet.write(i + 1, 13, fuzz_q1[i][12])
            elif len(fuzz_q1[i])==14:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
                worksheet.write(i + 1, 10, fuzz_q1[i][9])
                worksheet.write(i + 1, 11, fuzz_q1[i][10])
                worksheet.write(i + 1, 12, fuzz_q1[i][11])
                worksheet.write(i + 1, 13, fuzz_q1[i][12])
                worksheet.write(i + 1, 14, fuzz_q1[i][13])
            elif len(fuzz_q1[i])==15:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
                worksheet.write(i + 1, 10, fuzz_q1[i][9])
                worksheet.write(i + 1, 11, fuzz_q1[i][10])
                worksheet.write(i + 1, 12, fuzz_q1[i][11])
                worksheet.write(i + 1, 13, fuzz_q1[i][12])
                worksheet.write(i + 1, 14, fuzz_q1[i][13])
                worksheet.write(i + 1, 15, fuzz_q1[i][14])
            elif len(fuzz_q1[i])==16:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
                worksheet.write(i + 1, 10, fuzz_q1[i][9])
                worksheet.write(i + 1, 11, fuzz_q1[i][10])
                worksheet.write(i + 1, 12, fuzz_q1[i][11])
                worksheet.write(i + 1, 13, fuzz_q1[i][12])
                worksheet.write(i + 1, 14, fuzz_q1[i][13])
                worksheet.write(i + 1, 15, fuzz_q1[i][14])
                worksheet.write(i + 1, 16, fuzz_q1[i][15])
            elif len(fuzz_q1[i])==17:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
                worksheet.write(i + 1, 10, fuzz_q1[i][9])
                worksheet.write(i + 1, 11, fuzz_q1[i][10])
                worksheet.write(i + 1, 12, fuzz_q1[i][11])
                worksheet.write(i + 1, 13, fuzz_q1[i][12])
                worksheet.write(i + 1, 14, fuzz_q1[i][13])
                worksheet.write(i + 1, 15, fuzz_q1[i][14])
                worksheet.write(i + 1, 16, fuzz_q1[i][15])
                worksheet.write(i + 1, 17, fuzz_q1[i][16])
            elif len(fuzz_q1[i])==18:
                worksheet.write(i + 1, 9, fuzz_q1[i][8])
                worksheet.write(i + 1, 10, fuzz_q1[i][9])
                worksheet.write(i + 1, 11, fuzz_q1[i][10])
                worksheet.write(i + 1, 12, fuzz_q1[i][11])
                worksheet.write(i + 1, 13, fuzz_q1[i][12])
                worksheet.write(i + 1, 14, fuzz_q1[i][13])
                worksheet.write(i + 1, 15, fuzz_q1[i][14])
                worksheet.write(i + 1, 16, fuzz_q1[i][15])
                worksheet.write(i + 1, 17, fuzz_q1[i][16])
                worksheet.write(i + 1, 18, fuzz_q1[i][17])
            else:
                print("天气特征缺失")

        workbook.close()  # 关闭工作簿
    return render_template("fuzz_query.html")



if __name__ == '__main__':
    app.run(debug=True)

